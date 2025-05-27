import enum
import math
import random
import docx
# import fitz
from flask import Flask, render_template, request, jsonify, redirect, session
import psycopg2
from psycopg2.extras import RealDictCursor
import os
import hashlib
import re
import jwt
from datetime import datetime, timedelta
from functools import wraps
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import base64
import uuid
import pdfplumber
from werkzeug.utils import secure_filename
from langchain_community.vectorstores import Pinecone as PineconeVectorStore
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain.chains import RetrievalQA
from langchain_openai import ChatOpenAI
from PIL import Image
import pytesseract
import io
import uuid
from pinecone import Pinecone, ServerlessSpec
from dotenv import load_dotenv
import tempfile
import psycopg2
from psycopg2.extras import RealDictCursor
import jwt
from datetime import datetime, timedelta
from functools import wraps
from flask_cors import CORS
from psycopg2.extras import RealDictCursor
from sqlalchemy import create_engine, Column, Integer, String, Text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
import requests
import json
import re
from sqlalchemy.pool import QueuePool
from sqlalchemy import exc, text
import os
import tempfile
import shutil
import numpy as np
from sqlalchemy import func
from sklearn.metrics.pairwise import cosine_similarity
from flask import Flask, request, jsonify
from sqlalchemy import create_engine, Column, Integer, String, Text, Float, ARRAY, Enum, ForeignKey
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, relationship
from concurrent.futures import ThreadPoolExecutor
from openai import OpenAI
from flask_cors import CORS

load_dotenv()


app = Flask(__name__)
CORS(app)
app.config['SECRET_KEY'] = 'your_secret_key'

# Database configuration
# uaYoBkLLt6uu4Lgm
# DATABASE_URL = "postgresql://postgres:xpBS15qHiEWtm1pi@db.ciojbynixikujryugtvs.supabase.co:5432/postgres"
DATABASE_URL = "postgresql://postgres.ciojbynixikujryugtvs:uaYoBkLLt6uu4Lgm@aws-0-ap-south-1.pooler.supabase.com:5432/postgres"

# Email configuration
EMAIL_CONFIG = {
    'EMAIL_HOST': 'smtp.gmail.com',
    'PORT': 587,
    'EMAIL_USER': 'mohsinshahid052@gmail.com',
    'EMAIL_PASS': 'migknvusxsbamoqm'
}

UPLOAD_FOLDER = './uploads/profile_images'
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'jpg', 'jpeg', 'png'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


OPENAI_API_KEY = "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"
# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
print(OPENAI_API_KEY)
print(os.getenv("PINECONE_API_KEY"))
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_ENVIRONMENT = os.getenv("PINECONE_ENVIRONMENT")
JWT_SECRET = "your-jwt-secret-key"

pc = Pinecone(api_key="pcsk_5F1m4Q_54qKFj8oLheHRiSofwP2hsWaq7DG42gY4T3qs93eYs3sjmqZ6iJiuDoq8CgzRGU")

embeddings = OpenAIEmbeddings(
    model="text-embedding-3-large",
    openai_api_key=OPENAI_API_KEY
)

INDEX_NAME = "casemine"
dimension = 3072

if INDEX_NAME not in pc.list_indexes().names():
    pc.create_index(
        name=INDEX_NAME,
        dimension=dimension,
        metric="cosine",
        spec=ServerlessSpec(
            cloud='aws',
            region=PINECONE_ENVIRONMENT
        )
    )

index = pc.Index(INDEX_NAME)

vectorstore = PineconeVectorStore(
    index=index,
    embedding=embeddings,
    text_key="text",
    namespace="default"
)

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=100,
    length_function=len
)


def get_db_connection():
    conn = psycopg2.connect(DATABASE_URL)
    conn.autocommit = True
    return conn

def setup_database():
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id SERIAL PRIMARY KEY,
            first_name VARCHAR(100) NOT NULL,
            last_name VARCHAR(100) NOT NULL,
            country VARCHAR(100) NOT NULL,
            email VARCHAR(255) UNIQUE NOT NULL,
            phone_number VARCHAR(20) NOT NULL,
            password VARCHAR(255) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # Create blacklisted tokens table for logout functionality
    cur.execute('''
        CREATE TABLE IF NOT EXISTS blacklisted_tokens (
            id SERIAL PRIMARY KEY,
            token TEXT NOT NULL,
            blacklisted_on TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # Create contact messages table
    cur.execute('''
        CREATE TABLE IF NOT EXISTS contact_messages (
            id SERIAL PRIMARY KEY,
            subject VARCHAR(255) NOT NULL,
            message TEXT NOT NULL,
            email VARCHAR(255),
            user_id INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE SET NULL
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS profiles (
            id SERIAL PRIMARY KEY,
            user_id INTEGER UNIQUE NOT NULL,
            court_of_practice VARCHAR(255),
            organization VARCHAR(255),
            job_title VARCHAR(255),
            bio_graphy TEXT,
            profile_image_path VARCHAR(512),
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS cases (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL,
            case_name VARCHAR(255) NOT NULL,
            client_name VARCHAR(255) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
    ''')

    # Create folders table
    cur.execute('''
        CREATE TABLE IF NOT EXISTS folders (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL,
            folder_name VARCHAR(255) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS case_documents (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL,
            case_id INTEGER,
            folder_id INTEGER,
            document_name VARCHAR(255) NOT NULL,
            file_path VARCHAR(255) NOT NULL,
            file_type VARCHAR(100),
            file_size INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
            FOREIGN KEY (case_id) REFERENCES cases(id) ON DELETE CASCADE,
            FOREIGN KEY (folder_id) REFERENCES folders(id) ON DELETE CASCADE
    )
    ''')

    cur.execute('''
        CREATE TABLE IF NOT EXISTS attorney_profiles (
            id SERIAL PRIMARY KEY,
            user_id INTEGER UNIQUE NOT NULL,
            first_name VARCHAR(100) NOT NULL,
            last_name VARCHAR(100) NOT NULL,
            mobile_number VARCHAR(20) NOT NULL,
            show_mobile_number BOOLEAN DEFAULT FALSE,
            email VARCHAR(255) NOT NULL,
            show_email BOOLEAN DEFAULT FALSE,
            organization VARCHAR(255),
            website VARCHAR(255),
            total_experience FLOAT,
            consultation_fees FLOAT,
            job_title VARCHAR(255),
            education_degree VARCHAR(255),
            passing_year INTEGER,
            university_name VARCHAR(255),
            facebook_url VARCHAR(255),
            twitter_url VARCHAR(255),
            linkedin_url VARCHAR(255),
            address VARCHAR(255),
            city VARCHAR(100),
            province VARCHAR(100),
            postal_code VARCHAR(20),
            country VARCHAR(100),
            bio_graph TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
    ''')

    # Create table for practice areas
    cur.execute('''
        CREATE TABLE IF NOT EXISTS practice_areas (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL,
            area_name VARCHAR(255) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
    ''')

    # Create table for courts of practice
    cur.execute('''
        CREATE TABLE IF NOT EXISTS practice_courts (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL,
            court_name VARCHAR(255) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
    ''')

    # Create table for professional experience
    cur.execute('''
        CREATE TABLE IF NOT EXISTS professional_experience (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL,
            job_title VARCHAR(255) NOT NULL,
            company_name VARCHAR(255) NOT NULL,
            from_year INTEGER NOT NULL,
            to_year INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
    ''')

    # Create table for professional memberships
    cur.execute('''
        CREATE TABLE IF NOT EXISTS professional_memberships (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL,
            membership_name VARCHAR(255) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
    ''')

    # Create table for languages known
    cur.execute('''
        CREATE TABLE IF NOT EXISTS languages_known (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL,
            language_name VARCHAR(100) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
        )
    ''')


    conn.commit()
    cur.close()
    conn.close()

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def is_valid_email(email):
    pattern = r'^[\w\.-]+@[\w\.-]+\.\w+$'
    return re.match(pattern, email) is not None

def is_valid_phone(phone):
    pattern = r'^\+?[0-9]{10,15}$'
    return re.match(pattern, phone) is not None

def generate_token(user_id, email):
    payload = {
        'user_id': user_id,
        'email': email,
        'exp': datetime.utcnow() + timedelta(days=1)
    }
    return jwt.encode(payload, app.config['SECRET_KEY'], algorithm='HS256')

def send_email(subject, message, from_email=None):
    # Create message container
    msg = MIMEMultipart()
    msg['From'] = EMAIL_CONFIG['EMAIL_USER']
    msg['To'] = EMAIL_CONFIG['EMAIL_USER']  # Sending to yourself
    msg['Subject'] = f"Contact Form: {subject}"

    # Add sender info to the message if available
    body = message
    if from_email:
        body = f"From: {from_email}\n\n{message}"

    # Attach the message to the email
    msg.attach(MIMEText(body, 'plain'))

    try:
        # Setup the server
        server = smtplib.SMTP(EMAIL_CONFIG['EMAIL_HOST'], EMAIL_CONFIG['PORT'])
        server.starttls()  # Secure the connection
        server.login(EMAIL_CONFIG['EMAIL_USER'], EMAIL_CONFIG['EMAIL_PASS'])

        # Send the email
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        print(f"Email error: {e}")
        return False

def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None

        # Check if token is in headers
        if 'Authorization' in request.headers:
            auth_header = request.headers['Authorization']
            if auth_header.startswith('Bearer '):
                token = auth_header.split(' ')[1]

        if not token:
            return jsonify({'error': 'Token is missing'}), 401

        # Check if token is blacklisted
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT id FROM blacklisted_tokens WHERE token = %s", (token,))
        if cur.fetchone():
            cur.close()
            conn.close()
            return jsonify({'error': 'Token has been revoked'}), 401

        try:
            # Decode token
            data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
            user_id = data['user_id']

            # Get user info
            cur.execute("SELECT id, email FROM users WHERE id = %s", (user_id,))
            current_user = cur.fetchone()
            cur.close()
            conn.close()

            if not current_user:
                return jsonify({'error': 'User not found'}), 401

        except jwt.ExpiredSignatureError:
            return jsonify({'error': 'Token has expired'}), 401
        except jwt.InvalidTokenError:
            return jsonify({'error': 'Invalid token'}), 401

        return f(user_id, *args, **kwargs)

    return decorated

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def save_base64_image(base64_string, user_id):
    """Save a base64 image to the filesystem and return the path"""
    if not base64_string or ',' not in base64_string:
        return None

    # Extract the content type and the data
    format_info, base64_data = base64_string.split(',', 1)

    # Detect file extension from the format info
    if 'image/jpeg' in format_info:
        ext = 'jpg'
    elif 'image/png' in format_info:
        ext = 'png'
    elif 'image/gif' in format_info:
        ext = 'gif'
    else:
        # Default to jpg if we can't detect
        ext = 'jpg'

    try:
        # Decode the base64 data
        image_data = base64.b64decode(base64_data)

        # Generate a unique filename
        filename = f"profile_{user_id}_{uuid.uuid4().hex}.{ext}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        # Write the image to file
        with open(filepath, 'wb') as f:
            f.write(image_data)

        return filepath
    except Exception as e:
        print(f"Error saving image: {e}")
        return None

# setup_database()

@app.route('/api/signup', methods=['POST'])
def signup():
    data = request.json

    # Check if all required fields are present
    required_fields = ['first_name', 'last_name', 'country', 'email', 'phone_number', 'password']
    for field in required_fields:
        if field not in data:
            return jsonify({'error': f'{field} is required'}), 400

    # Validate email format
    if not is_valid_email(data['email']):
        return jsonify({'error': 'Invalid email format'}), 400

    # Validate phone number format
    if not is_valid_phone(data['phone_number']):
        return jsonify({'error': 'Invalid phone number format'}), 400

    # Check password length
    if len(data['password']) < 8:
        return jsonify({'error': 'Password must be at least 8 characters long'}), 400

    # Hash the password
    hashed_password = hash_password(data['password'])

    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Check if email already exists
        cur.execute("SELECT id FROM users WHERE email = %s", (data['email'],))
        if cur.fetchone():
            return jsonify({'error': 'Email already registered'}), 409

        # Insert new user
        cur.execute('''
            INSERT INTO users (first_name, last_name, country, email, phone_number, password)
            VALUES (%s, %s, %s, %s, %s, %s)
            RETURNING id, email
        ''', (
            data['first_name'],
            data['last_name'],
            data['country'],
            data['email'],
            data['phone_number'],
            hashed_password
        ))

        new_user = cur.fetchone()
        cur.close()
        conn.commit()
        conn.close()

        # Generate auth token
        token = generate_token(new_user['id'], new_user['email'])

        return jsonify({
            'message': 'User registered successfully',
            'user_id': new_user['id'],
            'email': new_user['email'],
            'token': token
        }), 201

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/login', methods=['POST'])
def login():
    data = request.json

    # Check if email and password are provided
    if 'email' not in data or 'password' not in data:
        return jsonify({'error': 'Email and password are required'}), 400

    hashed_password = hash_password(data['password'])

    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        cur.execute('''
            SELECT id, email, first_name, last_name
            FROM users
            WHERE email = %s AND password = %s
        ''', (data['email'], hashed_password))

        user = cur.fetchone()
        cur.close()
        conn.close()

        if not user:
            return jsonify({'error': 'Invalid email or password'}), 401

        # Generate auth token
        token = generate_token(user['id'], user['email'])

        return jsonify({
            'message': 'Login successful',
            'user': {
                'id': user['id'],
                'email': user['email'],
                'first_name': user['first_name'],
                'last_name': user['last_name']
            },
            'token': token
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/logout', methods=['POST'])
@token_required
def logout(user_id):
    # Get token from request header
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        return jsonify({'error': 'Invalid token format'}), 400

    token = auth_header.split(' ')[1]

    try:
        # Add token to blacklist
        conn = get_db_connection()
        cur = conn.cursor()

        cur.execute('''
            INSERT INTO blacklisted_tokens (token)
            VALUES (%s)
        ''', (token,))

        conn.commit()
        cur.close()
        conn.close()

        # Return success and redirect information
        return jsonify({
            'message': 'Successfully logged out',
            'redirect_url': '/'
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/change-password', methods=['POST'])
@token_required
def change_password(user_id):
    data = request.json

    # Check if current and new passwords are provided
    if 'current_password' not in data or 'new_password' not in data:
        return jsonify({'error': 'Current password and new password are required'}), 400

    # Validate new password length
    if len(data['new_password']) < 8:
        return jsonify({'error': 'New password must be at least 8 characters long'}), 400

    # Hash both passwords
    hashed_current_password = hash_password(data['current_password'])
    hashed_new_password = hash_password(data['new_password'])

    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Verify current password
        cur.execute('''
            SELECT id
            FROM users
            WHERE id = %s AND password = %s
        ''', (user_id, hashed_current_password))

        user = cur.fetchone()

        if not user:
            cur.close()
            conn.close()
            return jsonify({'error': 'Current password is incorrect'}), 401

        # Update password
        cur.execute('''
            UPDATE users
            SET password = %s
            WHERE id = %s
        ''', (hashed_new_password, user_id))

        conn.commit()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Password updated successfully'
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/contact', methods=['POST'])
def contact():
    data = request.json

    # Check if subject and message are provided
    if 'subject' not in data or 'message' not in data:
        return jsonify({'error': 'Subject and message are required'}), 400

    # Get email from authenticated user or from request
    user_id = None
    user_email = None

    # Check if user is authenticated
    auth_header = request.headers.get('Authorization')
    if auth_header and auth_header.startswith('Bearer '):
        token = auth_header.split(' ')[1]
        try:
            # Decode token without checking if it's blacklisted (we allow contact even from logged out users)
            token_data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
            user_id = token_data['user_id']
            user_email = token_data.get('email')
        except:
            pass  # Ignore token errors for contact form

    # Get email from request if not authenticated
    if not user_email and 'email' in data:
        user_email = data['email']

    try:
        # Save message to database
        conn = get_db_connection()
        cur = conn.cursor()

        cur.execute('''
            INSERT INTO contact_messages (subject, message, email, user_id)
            VALUES (%s, %s, %s, %s)
        ''', (data['subject'], data['message'], user_email, user_id))

        conn.commit()
        cur.close()
        conn.close()

        # Send email
        email_sent = send_email(data['subject'], data['message'], user_email)

        if email_sent:
            return jsonify({
                'message': 'Your message has been sent successfully!'
            }), 200
        else:
            return jsonify({
                'message': 'Your message was saved but there was an issue sending the email. We will still process your request.',
                'warning': 'Email delivery issue'
            }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route("/api/health", methods=['GET'])
def health():
    return jsonify({'status': 'ok'}), 200

@app.route('/api/profile', methods=['GET'])
@token_required
def get_profile(user_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Get user basic info
        cur.execute('''
            SELECT id, email, first_name, last_name, created_at
            FROM users
            WHERE id = %s
        ''', (user_id,))

        user_data = cur.fetchone()

        if not user_data:
            cur.close()
            conn.close()
            return jsonify({'error': 'User not found'}), 404

        # Get profile info
        cur.execute('''
            SELECT court_of_practice, organization, job_title,
                   bio_graphy, profile_image_path, updated_at
            FROM profiles
            WHERE user_id = %s
        ''', (user_id,))

        profile_data = cur.fetchone()

        cur.close()
        conn.close()

        # Combine user and profile data
        result = dict(user_data)

        if profile_data:
            result.update(profile_data)

            # Convert profile image path to URL if exists
            if profile_data['profile_image_path']:
                # In a real app, you'd use a proper URL based on your hosting
                result['profile_image_url'] = f"/api/images/{os.path.basename(profile_data['profile_image_path'])}"

        return jsonify(result), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/profile', methods=['PUT'])
@token_required
def update_profile(user_id):
    try:
        # Get form data from request
        # For a multi-part form with file uploads, use request.form and request.files
        # For a JSON request (without file uploads), use request.json

        form_data = {}

        # Check if content type is JSON
        if request.is_json:
            form_data = request.json
        else:
            # Handle form data
            form_data = request.form.to_dict()

        # Validate required fields
        required_fields = ['first_name', 'last_name']
        for field in required_fields:
            if field not in form_data or not form_data[field]:
                return jsonify({'error': f'{field} is required'}), 400

        # Optional fields with default empty values
        optional_fields = {
            'court_of_practice': '',
            'organization': '',
            'job_title': '',
            'bio_graphy': ''
        }

        # Set defaults for any missing optional fields
        for field, default in optional_fields.items():
            if field not in form_data:
                form_data[field] = default

        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Update user basic info
        cur.execute('''
            UPDATE users
            SET first_name = %s, last_name = %s
            WHERE id = %s
            RETURNING id, email, first_name, last_name
        ''', (
            form_data['first_name'],
            form_data['last_name'],
            user_id
        ))

        updated_user = cur.fetchone()

        # Check if profile exists
        cur.execute('SELECT id FROM profiles WHERE user_id = %s', (user_id,))
        profile_exists = cur.fetchone() is not None

        # Process profile image if present
        profile_image_path = None

        # Check for base64 image data
        if 'profile_image' in form_data and form_data['profile_image'] and form_data['profile_image'].startswith('data:image/'):
            profile_image_path = save_base64_image(form_data['profile_image'], user_id)

        # Check for file upload (multipart/form-data)
        elif request.files and 'profile_image' in request.files:
            file = request.files['profile_image']
            if file and file.filename and allowed_file(file.filename):
                filename = secure_filename(f"profile_{user_id}_{uuid.uuid4().hex}_{file.filename}")
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                profile_image_path = filepath

        # Update or insert profile
        if profile_exists:
            if profile_image_path:
                # Update with new image
                cur.execute('''
                    UPDATE profiles
                    SET court_of_practice = %s, organization = %s, job_title = %s,
                        bio_graphy = %s, profile_image_path = %s, updated_at = NOW()
                    WHERE user_id = %s
                    RETURNING id
                ''', (
                    form_data['court_of_practice'],
                    form_data['organization'],
                    form_data['job_title'],
                    form_data['bio_graphy'],
                    profile_image_path,
                    user_id
                ))
            else:
                # Update without changing image
                cur.execute('''
                    UPDATE profiles
                    SET court_of_practice = %s, organization = %s, job_title = %s,
                        bio_graphy = %s, updated_at = NOW()
                    WHERE user_id = %s
                    RETURNING id
                ''', (
                    form_data['court_of_practice'],
                    form_data['organization'],
                    form_data['job_title'],
                    form_data['bio_graphy'],
                    user_id
                ))
        else:
            # Insert new profile
            cur.execute('''
                INSERT INTO profiles (user_id, court_of_practice, organization, job_title,
                                     bio_graphy, profile_image_path)
                VALUES (%s, %s, %s, %s, %s, %s)
                RETURNING id
            ''', (
                user_id,
                form_data['court_of_practice'],
                form_data['organization'],
                form_data['job_title'],
                form_data['bio_graphy'],
                profile_image_path
            ))

        # Get the updated profile data
        cur.execute('''
            SELECT court_of_practice, organization, job_title,
                   bio_graphy, profile_image_path, updated_at
            FROM profiles
            WHERE user_id = %s
        ''', (user_id,))

        profile_data = cur.fetchone()

        conn.commit()
        cur.close()
        conn.close()

        # Combine user and profile data for response
        result = dict(updated_user)
        if profile_data:
            result.update(profile_data)

            # Convert profile image path to URL if exists
            if profile_data['profile_image_path']:
                # In a real app, you'd use a proper URL based on your hosting
                result['profile_image_url'] = f"/api/images/{os.path.basename(profile_data['profile_image_path'])}"

        return jsonify({
            'message': 'Profile updated successfully',
            'user': result
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/images/<filename>', methods=['GET'])
def get_image(filename):
    return app.send_from_directory(app.config['UPLOAD_FOLDER'], filename)

############################## Attorney IQ Api #################################################
@app.route('/api/attorney/profile', methods=['POST'])
@token_required
def create_attorney_profile(user_id):
    data = request.json

    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Check if profile already exists
        cur.execute("SELECT id FROM attorney_profiles WHERE user_id = %s", (user_id,))
        if cur.fetchone():
            return jsonify({'error': 'Profile already exists for this user. Use PUT method to update.'}), 409

        # Insert main profile data
        cur.execute('''
            INSERT INTO attorney_profiles (
                user_id, first_name, last_name, mobile_number, show_mobile_number,
                email, show_email, organization, website, total_experience,
                consultation_fees, job_title, education_degree, passing_year,
                university_name, facebook_url, twitter_url, linkedin_url,
                address, city, province, postal_code, country, bio_graph
            ) VALUES (
                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
            ) RETURNING id
        ''', (
            user_id,
            data.get('first_name', ''),
            data.get('last_name', ''),
            data.get('mobile_number', ''),
            data.get('show_mobile_number', False),
            data.get('email', ''),
            data.get('show_email', False),
            data.get('organization', ''),
            data.get('website', ''),
            data.get('total_experience', 0),
            data.get('consultation_fees', 0),
            data.get('job_title', ''),
            data.get('education_degree', ''),
            data.get('passing_year', None),
            data.get('university_name', ''),
            data.get('facebook_url', ''),
            data.get('twitter_url', ''),
            data.get('linkedin_url', ''),
            data.get('address', ''),
            data.get('city', ''),
            data.get('province', ''),
            data.get('postal_code', ''),
            data.get('country', ''),
            data.get('bio_graph', '')
        ))

        profile_id = cur.fetchone()['id']

        # Handle practice areas
        if 'practice_areas' in data and isinstance(data['practice_areas'], list):
            for area in data['practice_areas']:
                cur.execute(
                    "INSERT INTO practice_areas (user_id, area_name) VALUES (%s, %s)",
                    (user_id, area)
                )

        # Handle courts of practice
        if 'practice_courts' in data and isinstance(data['practice_courts'], list):
            for court in data['practice_courts']:
                cur.execute(
                    "INSERT INTO practice_courts (user_id, court_name) VALUES (%s, %s)",
                    (user_id, court)
                )

        # Handle professional experience
        if 'professional_experience' in data and isinstance(data['professional_experience'], list):
            for exp in data['professional_experience']:
                cur.execute(
                    "INSERT INTO professional_experience (user_id, job_title, company_name, from_year, to_year) VALUES (%s, %s, %s, %s, %s)",
                    (user_id, exp.get('job_title', ''), exp.get('company_name', ''),
                     exp.get('from_year', 0), exp.get('to_year'))
                )

        # Handle professional memberships
        if 'professional_memberships' in data and isinstance(data['professional_memberships'], list):
            for membership in data['professional_memberships']:
                cur.execute(
                    "INSERT INTO professional_memberships (user_id, membership_name) VALUES (%s, %s)",
                    (user_id, membership)
                )

        # Handle languages known
        if 'languages_known' in data and isinstance(data['languages_known'], list):
            for language in data['languages_known']:
                cur.execute(
                    "INSERT INTO languages_known (user_id, language_name) VALUES (%s, %s)",
                    (user_id, language)
                )

        conn.commit()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Attorney profile created successfully',
            'profile_id': profile_id
        }), 201

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/attorney/profile', methods=['PUT'])
@token_required
def update_attorney_profile(user_id):
    data = request.json

    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Check if profile exists
        cur.execute("SELECT id FROM attorney_profiles WHERE user_id = %s", (user_id,))
        profile = cur.fetchone()

        if not profile:
            return jsonify({'error': 'Profile not found. Create one first.'}), 404

        # Update main profile data
        cur.execute('''
            UPDATE attorney_profiles SET
                first_name = %s,
                last_name = %s,
                mobile_number = %s,
                show_mobile_number = %s,
                email = %s,
                show_email = %s,
                organization = %s,
                website = %s,
                total_experience = %s,
                consultation_fees = %s,
                job_title = %s,
                education_degree = %s,
                passing_year = %s,
                university_name = %s,
                facebook_url = %s,
                twitter_url = %s,
                linkedin_url = %s,
                address = %s,
                city = %s,
                province = %s,
                postal_code = %s,
                country = %s,
                bio_graph = %s,
                updated_at = CURRENT_TIMESTAMP
            WHERE user_id = %s
        ''', (
            data.get('first_name', ''),
            data.get('last_name', ''),
            data.get('mobile_number', ''),
            data.get('show_mobile_number', False),
            data.get('email', ''),
            data.get('show_email', False),
            data.get('organization', ''),
            data.get('website', ''),
            data.get('total_experience', 0),
            data.get('consultation_fees', 0),
            data.get('job_title', ''),
            data.get('education_degree', ''),
            data.get('passing_year', None),
            data.get('university_name', ''),
            data.get('facebook_url', ''),
            data.get('twitter_url', ''),
            data.get('linkedin_url', ''),
            data.get('address', ''),
            data.get('city', ''),
            data.get('province', ''),
            data.get('postal_code', ''),
            data.get('country', ''),
            data.get('bio_graph', ''),
            user_id
        ))

        # Handle practice areas - delete and recreate
        if 'practice_areas' in data and isinstance(data['practice_areas'], list):
            cur.execute("DELETE FROM practice_areas WHERE user_id = %s", (user_id,))
            for area in data['practice_areas']:
                cur.execute(
                    "INSERT INTO practice_areas (user_id, area_name) VALUES (%s, %s)",
                    (user_id, area)
                )

        # Handle courts of practice - delete and recreate
        if 'practice_courts' in data and isinstance(data['practice_courts'], list):
            cur.execute("DELETE FROM practice_courts WHERE user_id = %s", (user_id,))
            for court in data['practice_courts']:
                cur.execute(
                    "INSERT INTO practice_courts (user_id, court_name) VALUES (%s, %s)",
                    (user_id, court)
                )

        # Handle professional experience - delete and recreate
        if 'professional_experience' in data and isinstance(data['professional_experience'], list):
            cur.execute("DELETE FROM professional_experience WHERE user_id = %s", (user_id,))
            for exp in data['professional_experience']:
                cur.execute(
                    "INSERT INTO professional_experience (user_id, job_title, company_name, from_year, to_year) VALUES (%s, %s, %s, %s, %s)",
                    (user_id, exp.get('job_title', ''), exp.get('company_name', ''),
                     exp.get('from_year', 0), exp.get('to_year'))
                )

        # Handle professional memberships - delete and recreate
        if 'professional_memberships' in data and isinstance(data['professional_memberships'], list):
            cur.execute("DELETE FROM professional_memberships WHERE user_id = %s", (user_id,))
            for membership in data['professional_memberships']:
                cur.execute(
                    "INSERT INTO professional_memberships (user_id, membership_name) VALUES (%s, %s)",
                    (user_id, membership)
                )

        # Handle languages known - delete and recreate
        if 'languages_known' in data and isinstance(data['languages_known'], list):
            cur.execute("DELETE FROM languages_known WHERE user_id = %s", (user_id,))
            for language in data['languages_known']:
                cur.execute(
                    "INSERT INTO languages_known (user_id, language_name) VALUES (%s, %s)",
                    (user_id, language)
                )

        conn.commit()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Attorney profile updated successfully'
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/attorney/profile', methods=['GET'])
@token_required
def get_attorney_profile(user_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Get main profile data
        cur.execute("SELECT * FROM attorney_profiles WHERE user_id = %s", (user_id,))
        profile = cur.fetchone()

        if not profile:
            return jsonify({'error': 'Profile not found'}), 404

        # Get practice areas
        cur.execute("SELECT area_name FROM practice_areas WHERE user_id = %s", (user_id,))
        practice_areas = [row['area_name'] for row in cur.fetchall()]

        # Get courts of practice
        cur.execute("SELECT court_name FROM practice_courts WHERE user_id = %s", (user_id,))
        practice_courts = [row['court_name'] for row in cur.fetchall()]

        # Get professional experience
        cur.execute('''
            SELECT job_title, company_name, from_year, to_year
            FROM professional_experience
            WHERE user_id = %s
            ORDER BY from_year DESC
        ''', (user_id,))
        professional_experience = cur.fetchall()

        # Get professional memberships
        cur.execute("SELECT membership_name FROM professional_memberships WHERE user_id = %s", (user_id,))
        memberships = [row['membership_name'] for row in cur.fetchall()]

        # Get languages known
        cur.execute("SELECT language_name FROM languages_known WHERE user_id = %s", (user_id,))
        languages = [row['language_name'] for row in cur.fetchall()]

        # Build complete profile
        complete_profile = dict(profile)
        complete_profile['practice_areas'] = practice_areas
        complete_profile['practice_courts'] = practice_courts
        complete_profile['professional_experience'] = professional_experience
        complete_profile['professional_memberships'] = memberships
        complete_profile['languages_known'] = languages

        cur.close()
        conn.close()

        return jsonify(complete_profile), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/attorney/profile/<int:attorney_id>', methods=['GET'])
def get_public_attorney_profile(attorney_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        cur.execute('''
            SELECT
                ap.id, ap.first_name, ap.last_name, ap.organization, ap.job_title,
                ap.total_experience, ap.consultation_fees, ap.city, ap.country,
                ap.bio_graph, ap.website,
                CASE WHEN ap.show_email THEN ap.email ELSE NULL END as email,
                CASE WHEN ap.show_mobile_number THEN ap.mobile_number ELSE NULL END as mobile_number
            FROM attorney_profiles ap
            WHERE ap.user_id = %s
        ''', (attorney_id,))

        profile = cur.fetchone()

        if not profile:
            return jsonify({'error': 'Profile not found'}), 404

        cur.execute("SELECT area_name FROM practice_areas WHERE user_id = %s", (attorney_id,))
        practice_areas = [row['area_name'] for row in cur.fetchall()]

        cur.execute("SELECT court_name FROM practice_courts WHERE user_id = %s", (attorney_id,))
        practice_courts = [row['court_name'] for row in cur.fetchall()]

        cur.execute('''
            SELECT job_title, company_name, from_year, to_year
            FROM professional_experience
            WHERE user_id = %s
            ORDER BY from_year DESC
        ''', (attorney_id,))
        professional_experience = cur.fetchall()

        cur.execute("SELECT language_name FROM languages_known WHERE user_id = %s", (attorney_id,))
        languages = [row['language_name'] for row in cur.fetchall()]

        # Build complete profile for public view
        complete_profile = dict(profile)
        complete_profile['practice_areas'] = practice_areas
        complete_profile['practice_courts'] = practice_courts
        complete_profile['professional_experience'] = professional_experience
        complete_profile['languages_known'] = languages

        # Add social media links if available
        if profile.get('facebook_url'):
            complete_profile['facebook_url'] = profile['facebook_url']
        if profile.get('twitter_url'):
            complete_profile['twitter_url'] = profile['twitter_url']
        if profile.get('linkedin_url'):
            complete_profile['linkedin_url'] = profile['linkedin_url']

        cur.close()
        conn.close()

        return jsonify(complete_profile), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/attorney/search/document', methods=['POST'])
def search_attorneys_by_document():
    try:
        # Get the uploaded document and court name
        document = request.files.get('document')
        court_name = request.form.get('court_name', '')

        if not document:
            return jsonify({'error': 'No document provided'}), 400

        if not court_name:
            return jsonify({'error': 'Court name is required'}), 400

        # Process document (basic implementation - read text content)
        # In production, you'd use specific libraries for different file types
        try:
            document_text = document.read().decode('utf-8')
        except:
            document_text = ""  # If document can't be read as text

        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Find attorneys whose practice courts match the given court name
        query = """
            SELECT DISTINCT
                ap.user_id, ap.first_name, ap.last_name, ap.organization,
                ap.job_title, ap.city, ap.country, ap.total_experience
            FROM attorney_profiles ap
            JOIN practice_courts pc ON ap.user_id = pc.user_id
            WHERE LOWER(pc.court_name) LIKE LOWER(%s)
        """

        cur.execute(query, (f"%{court_name}%",))
        results = cur.fetchall()

        # Get additional details for each attorney
        for attorney in results:
            # Get practice areas
            cur.execute("SELECT area_name FROM practice_areas WHERE user_id = %s", (attorney['user_id'],))
            attorney['practice_areas'] = [row['area_name'] for row in cur.fetchall()]

            # Get courts of practice
            cur.execute("SELECT court_name FROM practice_courts WHERE user_id = %s", (attorney['user_id'],))
            attorney['practice_courts'] = [row['court_name'] for row in cur.fetchall()]

        cur.close()
        conn.close()

        return jsonify({
            'attorneys': results,
            'count': len(results),
            'search_criteria': {
                'court_name': court_name,
                'document_name': document.filename if document else None
            }
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/attorney/search/text', methods=['POST'])
def search_attorneys_by_text():
    try:
        # Get text input and court name
        data = request.json

        if not data:
            return jsonify({'error': 'No data provided'}), 400

        text = data.get('text', '')
        court_name = data.get('court_name', '')

        if not court_name:
            return jsonify({'error': 'Court name is required'}), 400

        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Find attorneys whose practice courts match the given court name
        query = """
            SELECT DISTINCT
                ap.user_id, ap.first_name, ap.last_name, ap.organization,
                ap.job_title, ap.city, ap.country, ap.total_experience
            FROM attorney_profiles ap
            JOIN practice_courts pc ON ap.user_id = pc.user_id
            WHERE LOWER(pc.court_name) LIKE LOWER(%s)
        """

        cur.execute(query, (f"%{court_name}%",))
        results = cur.fetchall()

        # Get additional details for each attorney
        for attorney in results:
            # Get practice areas
            cur.execute("SELECT area_name FROM practice_areas WHERE user_id = %s", (attorney['user_id'],))
            attorney['practice_areas'] = [row['area_name'] for row in cur.fetchall()]

            # Get courts of practice
            cur.execute("SELECT court_name FROM practice_courts WHERE user_id = %s", (attorney['user_id'],))
            attorney['practice_courts'] = [row['court_name'] for row in cur.fetchall()]

        cur.close()
        conn.close()

        return jsonify({
            'attorneys': results,
            'count': len(results),
            'search_criteria': {
                'court_name': court_name,
                'text_provided': bool(text)
            }
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

##########################################  Case and folder Api ##########################################

@app.route('/api/cases', methods=['POST'])
@token_required
def create_case(user_id):
    data = request.json

    # Check if required fields are present
    if 'case_name' not in data or 'client_name' not in data:
        return jsonify({'error': 'Case name and client name are required'}), 400

    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Insert new case
        cur.execute('''
            INSERT INTO cases (user_id, case_name, client_name)
            VALUES (%s, %s, %s)
            RETURNING id, case_name, client_name, created_at
        ''', (
            user_id,
            data['case_name'],
            data['client_name']
        ))

        new_case = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Case created successfully',
            'case': new_case
        }), 201

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/cases', methods=['GET'])
@token_required
def get_cases(user_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Get all cases for the user
        cur.execute('''
            SELECT id, case_name, client_name, created_at
            FROM cases
            WHERE user_id = %s
            ORDER BY created_at DESC
        ''', (user_id,))

        cases = cur.fetchall()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Cases retrieved successfully',
            'cases': cases
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/cases/<int:case_id>', methods=['PUT'])
@token_required
def update_case(user_id, case_id):
    data = request.json

    # Check if required fields are present
    if 'case_name' not in data and 'client_name' not in data:
        return jsonify({'error': 'At least one field to update is required'}), 400

    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Check if case exists and belongs to user
        cur.execute('''
            SELECT id FROM cases
            WHERE id = %s AND user_id = %s
        ''', (case_id, user_id))

        case = cur.fetchone()
        if not case:
            cur.close()
            conn.close()
            return jsonify({'error': 'Case not found or not authorized'}), 404

        # Build update query dynamically based on provided fields
        update_fields = []
        params = []

        if 'case_name' in data:
            update_fields.append("case_name = %s")
            params.append(data['case_name'])

        if 'client_name' in data:
            update_fields.append("client_name = %s")
            params.append(data['client_name'])

        # Complete the params and execute update
        params.append(case_id)
        params.append(user_id)

        query = f'''
            UPDATE cases
            SET {', '.join(update_fields)}
            WHERE id = %s AND user_id = %s
            RETURNING id, case_name, client_name, created_at
        '''

        cur.execute(query, params)
        updated_case = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Case updated successfully',
            'case': updated_case
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/cases/<int:case_id>', methods=['DELETE'])
@token_required
def delete_case(user_id, case_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Check if case exists and belongs to user
        cur.execute('''
            SELECT id FROM cases
            WHERE id = %s AND user_id = %s
        ''', (case_id, user_id))

        case = cur.fetchone()
        if not case:
            cur.close()
            conn.close()
            return jsonify({'error': 'Case not found or not authorized'}), 404

        # Delete the case
        cur.execute('''
            DELETE FROM cases
            WHERE id = %s AND user_id = %s
        ''', (case_id, user_id))

        conn.commit()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Case deleted successfully'
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/cases/<int:case_id>', methods=['GET'])
@token_required
def get_case(user_id, case_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Get case details
        cur.execute('''
            SELECT id, case_name, client_name, created_at
            FROM cases
            WHERE id = %s AND user_id = %s
        ''', (case_id, user_id))

        case = cur.fetchone()

        if not case:
            cur.close()
            conn.close()
            return jsonify({'error': 'Case not found or not authorized'}), 404

        # Get folders for this case
        cur.execute('''
            SELECT id, folder_name, created_at
            FROM folders
            WHERE case_id = %s
            ORDER BY created_at
        ''', (case_id,))

        folders = cur.fetchall()

        cur.close()
        conn.close()

        # Add folders to case data
        case_data = dict(case)
        case_data['folders'] = folders

        return jsonify({
            'message': 'Case retrieved successfully',
            'case': case_data
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/folders', methods=['POST'])
@token_required
def create_folder(user_id):
    data = request.json

    # Check if folder name is present
    if 'folder_name' not in data:
        return jsonify({'error': 'Folder name is required'}), 400

    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Insert new folder (with just a name and user_id)
        cur.execute('''
            INSERT INTO folders (user_id, folder_name)
            VALUES (%s, %s)
            RETURNING id, folder_name, created_at
        ''', (
            user_id,
            data['folder_name']
        ))

        new_folder = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Folder created successfully',
            'folder': new_folder
        }), 201

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/folders', methods=['GET'])
@token_required
def get_folders(user_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Get all folders for the user (without case relationship)
        cur.execute('''
            SELECT id, folder_name, created_at
            FROM folders
            WHERE user_id = %s
            ORDER BY created_at DESC
        ''', (user_id,))

        folders = cur.fetchall()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Folders retrieved successfully',
            'folders': folders
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/folders/<int:folder_id>', methods=['PUT'])
@token_required
def update_folder(user_id, folder_id):
    data = request.json

    # Check if folder name is present
    if 'folder_name' not in data:
        return jsonify({'error': 'Folder name is required'}), 400

    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Check if folder exists and belongs to user
        cur.execute('''
            SELECT id FROM folders
            WHERE id = %s AND user_id = %s
        ''', (folder_id, user_id))

        folder = cur.fetchone()
        if not folder:
            cur.close()
            conn.close()
            return jsonify({'error': 'Folder not found or not authorized'}), 404

        # Update the folder
        cur.execute('''
            UPDATE folders
            SET folder_name = %s
            WHERE id = %s AND user_id = %s
            RETURNING id, folder_name, created_at
        ''', (data['folder_name'], folder_id, user_id))

        updated_folder = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Folder updated successfully',
            'folder': updated_folder
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/folders/<int:folder_id>', methods=['DELETE'])
@token_required
def delete_folder(user_id, folder_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Check if folder exists and belongs to user
        cur.execute('''
            SELECT id FROM folders
            WHERE id = %s AND user_id = %s
        ''', (folder_id, user_id))

        folder = cur.fetchone()
        if not folder:
            cur.close()
            conn.close()
            return jsonify({'error': 'Folder not found or not authorized'}), 404

        # Delete the folder
        cur.execute('''
            DELETE FROM folders
            WHERE id = %s AND user_id = %s
        ''', (folder_id, user_id))

        conn.commit()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Folder deleted successfully'
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

############### document upload in folder API ###########################################

@app.route('/api/documents/upload', methods=['POST'])
@token_required
def upload_document(user_id):
    # Check if the post request has the file part
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']

    # If user does not select file, browser also submits an empty part without filename
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    # Check if it's an allowed file type
    if not allowed_file(file.filename):
        return jsonify({'error': f'File type not allowed. Allowed types: {", ".join(ALLOWED_EXTENSIONS)}'}), 400

    # Get parent container details (case or folder)
    parent_type = request.form.get('parent_type')  # 'case' or 'folder'
    parent_id = request.form.get('parent_id')  # ID of the case or folder

    if not parent_type or not parent_id or parent_type not in ['case', 'folder']:
        return jsonify({'error': 'Valid parent_type (case or folder) and parent_id are required'}), 400

    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Verify the parent container belongs to the user
        if parent_type == 'case':
            cur.execute('''
                SELECT id FROM cases
                WHERE id = %s AND user_id = %s
            ''', (parent_id, user_id))
        else:  # parent_type == 'folder'
            cur.execute('''
                SELECT id FROM folders
                WHERE id = %s AND user_id = %s
            ''', (parent_id, user_id))

        parent = cur.fetchone()
        if not parent:
            cur.close()
            conn.close()
            return jsonify({'error': f'{parent_type.capitalize()} not found or not authorized'}), 404

        # Save the file
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, f"{user_id}_{parent_type}_{parent_id}_{filename}")
        file.save(file_path)

        # Get file size
        file_size = os.path.getsize(file_path)

        # Get file type
        file_type = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

        # Insert document record in database
        case_id = parent_id if parent_type == 'case' else None
        folder_id = parent_id if parent_type == 'folder' else None

        cur.execute('''
            INSERT INTO case_documents
            (user_id, case_id, folder_id, document_name, file_path, file_type, file_size)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            RETURNING id, document_name, file_path, file_type, file_size, created_at
        ''', (
            user_id,
            case_id,
            folder_id,
            filename,
            file_path,
            file_type,
            file_size
        ))

        document = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Document uploaded successfully',
            'document': document
        }), 201

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/documents', methods=['GET'])
@token_required
def get_documents(user_id):
    # Get query parameters
    parent_type = request.args.get('parent_type')  # 'case' or 'folder'
    parent_id = request.args.get('parent_id')

    if not parent_type or not parent_id or parent_type not in ['case', 'folder']:
        return jsonify({'error': 'Valid parent_type (case or folder) and parent_id are required'}), 400

    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Verify the parent container belongs to the user
        if parent_type == 'case':
            cur.execute('''
                SELECT id FROM cases
                WHERE id = %s AND user_id = %s
            ''', (parent_id, user_id))

            if cur.fetchone():
                # Get documents for this case
                cur.execute('''
                    SELECT id, document_name, file_path, file_type, file_size, created_at
                    FROM case_documents
                    WHERE case_id = %s AND user_id = %s
                    ORDER BY created_at DESC
                ''', (parent_id, user_id))
            else:
                cur.close()
                conn.close()
                return jsonify({'error': 'Case not found or not authorized'}), 404

        else:  # parent_type == 'folder'
            cur.execute('''
                SELECT id FROM folders
                WHERE id = %s AND user_id = %s
            ''', (parent_id, user_id))

            if cur.fetchone():
                # Get documents for this folder
                cur.execute('''
                    SELECT id, document_name, file_path, file_type, file_size, created_at
                    FROM case_documents
                    WHERE folder_id = %s AND user_id = %s
                    ORDER BY created_at DESC
                ''', (parent_id, user_id))
            else:
                cur.close()
                conn.close()
                return jsonify({'error': 'Folder not found or not authorized'}), 404

        documents = cur.fetchall()
        cur.close()
        conn.close()

        return jsonify({
            'message': 'Documents retrieved successfully',
            'documents': documents
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500


#################### AMICUS ############################################################

def save_chat_message(user_id, conversation_id, message, response, document_id=None):
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor()
            cur.execute(
                "INSERT INTO amicus (conversation_id, user_id, message, response, document_id) VALUES (%s, %s, %s, %s, %s)",
                (conversation_id, user_id, message, response, document_id)
            )
            conn.commit()
            cur.close()
            conn.close()
            return True
        except Exception as e:
            print(f"Error saving chat message: {e}")
            return False
    return False

def get_conversation_history(conversation_id, user_id):
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute(
                "SELECT * FROM amicus WHERE conversation_id = %s AND user_id = %s ORDER BY created_at ASC",
                (conversation_id, str(user_id))
            )
            history = cur.fetchall()
            cur.close()
            conn.close()
            return history
        except Exception as e:
            print(f"Error retrieving conversation history: {e}")
            return []
        return []

def extract_text_from_pdf(pdf_file):
    """Extract text from a PDF using LangChain's loader"""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
        tmp_file.write(pdf_file.read())
        tmp_path = tmp_file.name

    try:
        loader = PyPDFLoader(tmp_path)
        documents = loader.load()
        return "\n".join([doc.page_content for doc in documents])
    finally:
        os.unlink(tmp_path)

def extract_text_from_image(image_file):
    """Extract text from an image using OCR"""
    try:
        image = Image.open(image_file)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        raise Exception(f"Image processing failed: {str(e)}")

def process_file(file):
    """Process uploaded file, extract text, create embeddings, and store in Pinecone"""
    try:
        file_content = file.read()
        file_io = io.BytesIO(file_content)
        file_type = file.filename.split('.')[-1].lower()

        if file_type == 'pdf':
            file_io.seek(0)
            text = extract_text_from_pdf(file_io)
        elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
            text = extract_text_from_image(file_io)
        else:
            return {"error": f"Unsupported file type: {file_type}"}

        docs = text_splitter.create_documents([text])

        doc_id = str(uuid.uuid4())
        for i, doc in enumerate(docs):
            doc.metadata = {
                "doc_id": doc_id,
                "chunk_id": i,
                "filename": file.filename
            }

        vectorstore.add_documents(docs)

        return {
            "success": True,
            "doc_id": doc_id,
            "chunks_processed": len(docs),
        }
    except Exception as e:
        return {"error": f"File processing failed: {str(e)}"}

def query_documents(query_text, top_k=5):
    """Query Pinecone for similar documents"""
    try:
        retriever = vectorstore.as_retriever(search_kwargs={"k": top_k})
        docs = retriever.get_relevant_documents(query_text)

        results = []
        for i, doc in enumerate(docs):
            results.append({
                "id": i,
                "text": doc.page_content,
                "metadata": doc.metadata,
                "score": i
            })
        return {"results": results}
    except Exception as e:
        raise Exception(f"Query failed: {str(e)}")

# New API route to create a new conversation
@app.route('/conversations/new', methods=['POST'])
@token_required
def create_conversation(user_id):
    conversation_id = str(uuid.uuid4())
    return jsonify({
        "conversation_id": conversation_id,
        "created_at": datetime.now().isoformat(),
        "user_id": user_id
    })

# API route to get conversation history
@app.route('/conversations/<conversation_id>/history', methods=['GET'])
@token_required
def get_history(user_id, conversation_id):
    history = get_conversation_history(conversation_id, user_id)
    return jsonify({"history": history})

# @app.route('/upload', methods=['POST'])
# @token_required
# def upload_file(user_id):
#     if 'file' not in request.files:
#         return jsonify({"error": "No file part"}), 400

#     file = request.files['file']
#     if file.filename == '':
#         return jsonify({"error": "No selected file"}), 400

#     # Get conversation_id from request
#     conversation_id = request.form.get('conversation_id')
#     if not conversation_id:
#         return jsonify({"error": "No conversation ID provided"}), 400

#     result = process_file(file)
#     if "error" in result:
#         return jsonify(result), 400

#     # Store the doc_id in the session
#     session['last_doc_id'] = result['doc_id']

#     # Save file upload action to chat history
#     message = f"Uploaded file: {file.filename}"
#     response = f"File processed successfully. Document ID: {result['doc_id']}"
#     save_chat_message(user_id, conversation_id, message, response, result['doc_id'])

#     return jsonify(result)

@app.route('/query', methods=['POST'])
@token_required
def query(user_id):
    data = request.json
    if not data or 'query' not in data:
        return jsonify({"error": "No query provided"}), 400

    conversation_id = data.get('conversation_id')
    if not conversation_id:
        return jsonify({"error": "No conversation ID provided"}), 400

    results = query_documents(data['query'], top_k=data.get('top_k', 5))

    # Save query to chat history
    save_chat_message(user_id, conversation_id, data['query'], str(results))

    return jsonify(results)

@app.route('/qa', methods=['POST'])
@token_required
def question_answer(user_id):
    data = request.json
    if not data or 'question' not in data:
        return jsonify({"error": "No question provided"}), 400

    conversation_id = data.get('conversation_id')
    if not conversation_id:
        return jsonify({"error": "No conversation ID provided"}), 400

    try:
        llm = ChatOpenAI(
            openai_api_key=OPENAI_API_KEY,
            model_name="gpt-4o"
        )

        retriever = vectorstore.as_retriever(search_kwargs={"k": 5})

        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=retriever
        )

        prompt = f"""
        You are a legal expert specializing in Indian law. Answer the following question in the context of Indian law, specifically referencing the Indian Contract Act, 1872, and other relevant statutes where applicable. Provide a detailed explanation, including:

        1. The definition or explanation of the legal concept in question as per Indian law.
        2. Relevant sections from the Indian Contract Act, 1872, or other statutes that apply.
        3. Landmark judgments from the Indian High Courts or Supreme Court that illustrate the application of the law, including the case name, court, and year.
        4. A reasoning section that ties the legal provisions and case law to the question.

        Ensure the answer is concise, accurate, and formatted clearly with the above sections. If the question is not directly related to Indian law, provide a general answer and note that it may not fully apply in the Indian legal context.

        Question: {data['question']}
        """

        answer = qa_chain.invoke(prompt)
        result = answer['result']

        # Save Q&A to chat history
        document_id = session.get('last_doc_id')
        save_chat_message(user_id, conversation_id, data['question'], result, document_id)

        return jsonify({
            "question": data['question'],
            "answer": result
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/summarize', methods=['POST'])
@token_required
def summarize(user_id):
    try:
        data = request.json
        conversation_id = data.get('conversation_id')
        doc_id = data.get("last_doc_id")
        if not conversation_id:
            return jsonify({"error": "No conversation ID provided"}), 400

        if  not doc_id:
            return jsonify({"error": "No document uploaded. Please upload a document first."}), 400

        retriever = vectorstore.as_retriever(search_kwargs={"k": 1, "filter": {"doc_id": doc_id}})
        docs = retriever.get_relevant_documents("summarize documents")
        if not docs:
            return jsonify({"error": "No documents found to summarize"}), 400

        document_title = docs[0].metadata.get("filename", "Document")

        llm = ChatOpenAI(
            openai_api_key=OPENAI_API_KEY,
            model_name="gpt-4o"
        )
        retriever = vectorstore.as_retriever(search_kwargs={"k": 5, "filter": {"doc_id": doc_id}})
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=retriever
        )

        prompt = f"""
        Provide a detailed summary of the uploaded documents in the following structured format. Ensure that each section is clearly labeled and contains relevant information based on the document's content. If certain information is not available, make a reasonable inference or note that the information is not specified.

        **Summary of {document_title}**

        1. **Type of Document**
        Identify the type of document (e.g., Project Report, Contract, Research Paper) and provide a brief description of its nature.

        2. **Purpose and Scope**
        Explain the purpose of the document and its scope. Include the objectives, what the document aims to achieve, and any limitations or constraints mentioned.

        3. **Key Parties**
        List the primary parties involved in the document (e.g., client, author, team, contractors) and describe their roles.

        4. **Main Provisions/Clauses**
        Highlight the main sections, provisions, or clauses in the document. Summarize their content briefly.

        5. **Legal References and Authority**
        Identify any legal references, regulations, or standards cited in the document. Explain their relevance to the document's validity or enforceability.

        6. **Obligations or Duties**
        Outline the obligations or duties of the involved parties as specified in the document.

        7. **Rights and Protections**
        Describe the rights and protections afforded to the parties involved, such as approval rights, safety guarantees, or compliance with standards.

        8. **Remedies or Penalties**
        Detail any remedies or penalties mentioned in the document for noncompliance or breaches of obligations.

        9. **Key Dates and Deadlines**
        Mention any specific dates, deadlines, or timelines outlined in the document. If none are specified, note that timelines may need to be established.

        10. **Limitations or Exclusions**
        Identify any limitations, exclusions, or constraints mentioned in the document that may affect its implementation or outcomes.

        11. **Overall Impact**
        Discuss the broader significance of the document, such as its potential impact on the parties involved, the project, or the field.

        12. **Additional Observations**
        Provide any additional observations or notable points about the document, such as unique features, automation, or potential for modifications.

        Ensure the summary is concise yet informative, with each section clearly separated and labeled as shown above.
        """

        summary = qa_chain.invoke(prompt)
        result = summary['result']

        # Save summarization to chat history
        message = f"Request to summarize document: {document_title}"
        save_chat_message(user_id, conversation_id, message, result, doc_id)

        return jsonify({"summary": result})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/argument', methods=['POST'])
@token_required
def argument(user_id):
    data = request.json
    if not data or 'query' not in data:
        return jsonify({"error": "No query provided"}), 400

    conversation_id = data.get('conversation_id')
    if not conversation_id:
        return jsonify({"error": "No conversation ID provided"}), 400

    argument_type = data.get('type', 'both')

    try:
        llm = ChatOpenAI(
            openai_api_key=OPENAI_API_KEY,
            model_name="gpt-4o"
        )
        retriever = vectorstore.as_retriever(search_kwargs={"k": 5})
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=retriever
        )

        # Different prompts based on argument type
        if argument_type == 'for':
            prompt = f"""
            You are a legal expert specializing in Indian law. Generate comprehensive legal arguments IN FAVOR of the following case or legal position, focusing exclusively on supporting arguments:

            Case: {data['query']}

            Please structure your response as follows:

            **Arguments In Favor**

            1. **Legal Basis**
               - Cite specific provisions from relevant Indian statutes (with section numbers) that support this position
               - Explain how these provisions directly apply to the case

            2. **Landmark Judgments**
               - Reference at least 3 landmark judgments from Indian courts (including case name, court, and year) that support this position
               - Explain the precedent established in each case and how it applies

            3. **Constitutional Principles**
               - Identify any constitutional provisions or principles that strengthen this position
               - Explain their relevance to the case

            4. **Legal Reasoning**
               - Provide compelling legal reasoning that connects the statutes, judgments, and principles to the case
               - Address potential counterarguments briefly only to demonstrate how they can be overcome

            5. **Summary of Position**
               - Conclude with a concise summary of why this position is legally sound under Indian law

            Focus only on arguments that SUPPORT the position/case, with emphasis on Indian legal framework. Be thorough but concise in your analysis.
            """
        elif argument_type == 'against':
            prompt = f"""
            You are a legal expert specializing in Indian law. Generate comprehensive legal arguments AGAINST the following case or legal position, focusing exclusively on opposing arguments:

            Case: {data['query']}

            Please structure your response as follows:

            **Arguments Against**

            1. **Legal Barriers**
               - Cite specific provisions from relevant Indian statutes (with section numbers) that oppose this position
               - Explain how these provisions create legal obstacles for the case

            2. **Contradicting Judgments**
               - Reference at least 3 landmark judgments from Indian courts (including case name, court, and year) that contradict this position
               - Explain how these precedents undermine the case

            3. **Constitutional Challenges**
               - Identify any constitutional provisions or principles that weaken this position
               - Explain their relevance to the case

            4. **Legal Reasoning**
               - Provide compelling legal reasoning that demonstrates why this position faces significant legal challenges
               - Highlight logical flaws or legal inconsistencies in the position

            5. **Summary of Opposition**
               - Conclude with a concise summary of why this position may not be legally sustainable under Indian law

            Focus only on arguments that OPPOSE the position/case, with emphasis on Indian legal framework. Be thorough but concise in your analysis.
            """
        else:  # 'both' - provide balanced arguments
            prompt = f"""
            You are a legal expert specializing in Indian law. Generate comprehensive legal arguments both FOR and AGAINST the following case or legal position:

            Case: {data['query']}

            Please structure your response with two clearly separated sections:

            **Arguments In Favor**

            1. **Legal Basis**
               - Cite specific provisions from relevant Indian statutes (with section numbers) that support this position
               - Explain how these provisions directly apply to the case

            2. **Supporting Judgments**
               - Reference at least 2 landmark judgments from Indian courts (including case name, court, and year) that support this position
               - Explain the precedent established in each case and how it applies

            3. **Legal Reasoning**
               - Provide compelling legal reasoning that connects the statutes, judgments, and principles to the case

            **Arguments Against**

            1. **Legal Barriers**
               - Cite specific provisions from relevant Indian statutes (with section numbers) that oppose this position
               - Explain how these provisions create legal obstacles for the case

            2. **Contradicting Judgments**
               - Reference at least 2 landmark judgments from Indian courts (including case name, court, and year) that contradict this position
               - Explain how these precedents undermine the case

            3. **Legal Reasoning**
               - Provide compelling legal reasoning that demonstrates why this position faces significant legal challenges

            **Balanced Conclusion**
            Provide a brief, balanced assessment of the legal position based on Indian law, considering the strength of arguments on both sides.

            Ensure thorough analysis with emphasis on the Indian legal framework. Be concise but comprehensive in your analysis.
            """

        arguments = qa_chain.invoke(prompt)
        result = arguments['result']

        # Save arguments to chat history
        document_id = session.get('last_doc_id')
        save_chat_message(user_id, conversation_id, data['query'], result, document_id)

        return jsonify({"arguments": result})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/contract_review', methods=['POST'])
@token_required
def contract_review(user_id):
    try:
        data = request.json
        conversation_id = data.get('conversation_id')
        doc_id = data.get('last_doc_id')
        if not conversation_id:
            return jsonify({"error": "No conversation ID provided"}), 400

        if not doc_id:
            return jsonify({"error": "No document uploaded. Please upload a document first."}), 400


        llm = ChatOpenAI(
            openai_api_key=OPENAI_API_KEY,
            model_name="gpt-4o"
        )

        retriever = vectorstore.as_retriever(search_kwargs={"k": 5, "filter": {"doc_id": doc_id}})

        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=retriever
        )

        prompt = """
        You are a legal expert specializing in Indian law. Review the uploaded document as if it were a contract and provide a detailed analysis, focusing on its compliance with Indian law, specifically the Indian Contract Act, 1872. Structure your response as follows:

        **Contract Review Report**

        1. **Document Type and Overview**
        Identify whether this document is a contract or another type of legal document. If it's not a contract, briefly explain what type of document it appears to be, then proceed to analyze it as if it were a contract or similar legal document.
        Provide a brief summary of the document, including its purpose, parties involved (if any), and main objectives.

        2. **Key Clauses or Sections**
        Identify and summarize the key clauses or sections in the document.

        3. **Compliance with Indian Law**
        Analyze whether the document complies with Indian law, specifically the Indian Contract Act, 1872, and other relevant legislation. Consider the following elements as applicable:
        - Offer and acceptance (Section 2(a) and 2(b))
        - Lawful consideration (Section 10 and Section 23)
        - Competent parties (Section 11)
        - Free consent (Sections 13-19)
        - Lawful object (Section 23)
        Cite the relevant sections of the Indian Contract Act, 1872, and explain how the document adheres to or violates these provisions.

        4. **Landmark Judgments**
        Reference at least two relevant landmark judgments from the Indian Supreme Court or High Courts that relate to the document's compliance or issues. Include the case name, court, and year, and explain how these judgments apply to the document.

        5. **Potential Issues**
        Highlight any potential legal issues or risks in the document, such as ambiguous language, non-compliance with Indian law, or sections that might be unenforceable.

        6. **Recommendations**
        Provide recommendations to address the identified issues, such as revising specific sections, ensuring compliance with Indian law, or adding protective measures.

        7. **Compliance Verdict**
        State whether the document is compliant with Indian law based on your analysis. If it is not compliant, specify the reasons and suggest corrective actions.

        Ensure the response is concise, accurate, and formatted clearly with the above sections. Use the content of the uploaded document as the primary source for the review.
        """

        review = qa_chain.invoke(prompt)
        result = review['result']

        message = "Request for contract review"
        save_chat_message(user_id, conversation_id, message, result, doc_id)

        return jsonify({"review": result})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/draft', methods=['POST'])
@token_required
def draft(user_id):
    data = request.json
    if not data or 'prompt' not in data:
        return jsonify({"error": "No prompt provided"}), 400

    conversation_id = data.get('conversation_id')
    if not conversation_id:
        return jsonify({"error": "No conversation ID provided"}), 400

    try:
        # Use only the prompt, skip document context
        prompt_lower = data['prompt'].lower()
        contract_types = {
            "nda": "Non-Disclosure Agreement (NDA)",
            "legal contract": "General Legal Contract",
            "service agreement": "Service Agreement",
            "partnership agreement": "Partnership Agreement",
            "employment contract": "Employment Contract"
        }

        contract_name = None
        for key, value in contract_types.items():
            if key in prompt_lower:
                contract_name = value
                break

        if not contract_name:
            return jsonify({
                "error": "Please specify a valid contract type in the prompt (e.g., 'Draft an NDA', 'Draft a Legal Contract'). Supported types: NDA, Legal Contract, Service Agreement, Partnership Agreement, Employment Contract."
            }), 400

        company_a = "Company A"
        company_b = "Company B"
        if "between" in prompt_lower:
            try:
                parts = data['prompt'].split("between")[1].split("and")
                company_a = parts[0].strip()
                company_b = parts[1].strip()
            except:
                pass

        draft_prompt = f"""
        You are a legal expert specializing in drafting contracts. Draft a professional {contract_name} between two companies, {company_a} and {company_b}, incorporated under the laws of India. Ensure the contract is a standalone, valid document. Structure the contract as follows:

        *{contract_name.upper()}*

        *Introduction*
        - State that this is a {contract_name} entered into as of [Effective Date] between {company_a} and {company_b}.
        - Provide placeholders for company addresses (e.g., [Address of {company_a}], [Address of {company_b}]).
        - Define the parties collectively as "Parties" and individually as a "Party."

        *Clauses*
        Include the following standard clauses, tailored to the {contract_name}:

        1. *Purpose*
        Define the purpose of the contract.

        2. *Definitions*
        Define key terms relevant to the contract.

        3. *Obligations of the Parties*
        Outline the obligations of both {company_a} and {company_b}.

        4. *Term and Termination*
        Specify the duration of the contract and conditions for termination.

        5. *Remedies*
        Include a clause on remedies for breach.

        6. *Governing Law*
        State that the contract is governed by the laws of India.

        7. *Dispute Resolution*
        Include a clause for dispute resolution, such as arbitration in India.

        8. *Miscellaneous*
        Add standard miscellaneous clauses.

        *Formatting*
        - Use a professional format with numbered sections and clear headings.
        - Include placeholders (e.g., [Effective Date], [Address of {company_a}]).
        - Add a disclaimer at the top stating that this is a sample contract for informational purposes and should be reviewed by a legal professional before execution.
        """

        llm = ChatOpenAI(
            openai_api_key=OPENAI_API_KEY,
            model_name="gpt-4o"
        )
        response = llm.invoke(draft_prompt)
        return jsonify({"draft": response.content})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


########################### CASE IQ ####################################################


conn = None
engine = None
Session = None
Base = declarative_base()

class Document(Base):
    __tablename__ = 'documents'
    id = Column(Integer, primary_key=True)
    filename = Column(String(255), nullable=False)
    text = Column(Text, nullable=False)
    facts = Column(Text)
    issues = Column(Text)
    petitioner_arguments = Column(Text)
    respondent_arguments = Column(Text)
    holding = Column(Text)
    reasoning = Column(Text)

class CourtType(enum.Enum):
    SUPREME_COURT = "supreme_court"
    HIGH_COURT = "high_court"
    TRIBUNAL = "tribunal"

class Article(Base):
    __tablename__ = 'articles'
    id = Column(Integer, primary_key=True)
    name = Column(String(255), nullable=False)
    text = Column(Text, nullable=False)

class Article2(Base):
    __tablename__ = 'article2'
    id = Column(Integer, primary_key=True)
    name = Column(String(255), nullable=False)
    text = Column(Text, nullable=False)
    year = Column(Integer, nullable=True)
    court_type = Column(Enum(CourtType), nullable=False)
    embedding = Column(ARRAY(Float), nullable=True)
    chunks = relationship("ArticleChunk", back_populates="article")

class ArticleChunk(Base):
    __tablename__ = 'article_chunks'
    id = Column(Integer, primary_key=True)
    article_id = Column(Integer, ForeignKey('article2.id'), nullable=False)
    chunk_number = Column(Integer, nullable=False)
    chunk_text = Column(Text, nullable=False)
    embedding = Column(ARRAY(Float), nullable=True)
    article = relationship("Article2", back_populates="chunks")


client = OpenAI(api_key="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")


def get_db_session():
    global Session, db_connected
    if not db_connected:
        db_connected = connect_to_db()

    if db_connected and Session is not None:
        try:
            session = Session()
            # Test the session with a simple query
            session.execute(text("SELECT 1"))
            return session
        except Exception as e:
            print(f"Session error: {e}")
            Session.remove()  # Clear any bad sessions
            db_connected = connect_to_db()  # Reconnect
            if db_connected:
                return Session()
    return None


def connect_to_db():
    global conn, engine, Session
    try:
        # Properly configure connection pooling to prevent pool exhaustion
        engine = create_engine(
            DATABASE_URL,
            poolclass=QueuePool,
            pool_size=5,       # Reduced pool size
            max_overflow=10,   # Limited overflow connections
            pool_timeout=30,
            pool_recycle=300,  # Recycle connections after 5 minutes
            pool_pre_ping=True # Verify connections before use
        )

        conn = psycopg2.connect(DATABASE_URL)
        Base.metadata.create_all(engine)

        # Use scoped session to better handle session lifecycle
        from sqlalchemy.orm import scoped_session
        Session = scoped_session(sessionmaker(bind=engine))

        print("Connected to Neon PostgreSQL database!")
        return True
    except Exception as e:
        print(f"Database connection error: {e}")
        return False


db_connected = connect_to_db()

def extract_text(file_path):
    try:
        if file_path.endswith('.pdf'):
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ''  # Prevent NoneType issues
            return text
        elif file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            return ""
    except Exception as e:
        print(f"Error extracting text: {e}")
        return "Error extracting text from document."
@app.route('/create-tables', methods=['GET'])
def create_tables():
    try:
        if not db_connected:
            connect_to_db()

        Base.metadata.create_all(engine)

        cursor = conn.cursor()
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id SERIAL PRIMARY KEY,
            filename VARCHAR(255) NOT NULL,
            text TEXT NOT NULL,
            facts TEXT,
            issues TEXT,
            petitioner_arguments TEXT,
            respondent_arguments TEXT,
            holding TEXT,
            reasoning TEXT
        )
        """)
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS articles (
            id SERIAL PRIMARY KEY,
            name VARCHAR(255) NOT NULL,
            text TEXT NOT NULL
        )
        """)
        conn.commit()
        cursor.close()

        return jsonify({
            'status': 'Success',
            'message': 'Tables created successfully'
        })
    except Exception as e:
        return jsonify({
            'status': 'Error',
            'message': f'Failed to create tables: {str(e)}'
        })


def analyze_legal_document(text):
    if not text or len(text.strip()) < 50:
        return {
            "facts": "Insufficient text to analyze.",
            "issues": "Insufficient text to analyze.",
            "petitioner_arguments": "Insufficient text to analyze.",
            "respondent_arguments": "Insufficient text to analyze.",
            "holding": "Insufficient text to analyze.",
            "reasoning": "Insufficient text to analyze."
        }

    try:
        max_text_length = 6000
        truncated_text = text[:max_text_length] if len(text) > max_text_length else text

        prompt = f"""
        Analyze the following text and extract the specific components below in the required format.
        If the document is not a legal case, still extract information following these categories.

        # DOCUMENT TEXT:
        {truncated_text}

        # FORMAT YOUR RESPONSE EXACTLY AS FOLLOWS WITH THESE EXACT HEADINGS:

        **facts**
        [Extract all relevant factual information, technological details, and background information. Include what the system does, its features, and capabilities.]

        **issues**
        [Identify key questions or concerns raised by the document. For technology documents, focus on potential challenges, implications, or considerations of using the described system.]

        **petitionerArguments**
        [Extract or infer positive points, benefits, or advantages mentioned or implied in the document.]

        **respondentArguments**
        [Extract or infer potential concerns, disadvantages, or counterarguments that might be raised against what's described in the document.]

        **holding**
        [For legal cases: the court's decision. For other documents: a summary conclusion or main takeaway from the document.]

        **reasoning**
        [Explain the logic, rationale, or justification behind the main points in the document. For technology documents, explain why certain features exist and their implications.]

        Do not include the quotes around the section titles and do not include any notes about the document not being a legal case. Just fill in each section with appropriate content from the document.
        Keep your extraction focused and make sure you don't repeat the same information across sections.
        """

        system_message = """
        You are an expert document analyzer that extracts structured information from texts.
        Your task is to extract key components from documents with precision and accuracy.
        Format your response using the exact section headings requested without adding any explanatory text.
        For each section, provide concise but comprehensive content relevant to that category.
        Do NOT include phrases like "Not provided in the document" or explanations about the document type.
        Instead, synthesize available information to fill each requested section appropriately.
        """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.1
        )

        full_response = response.choices[0].message.content
        print("Analysis completed - processing response")

        analysis = {
            "facts": "",
            "issues": "",
            "petitioner_arguments": "",
            "respondent_arguments": "",
            "holding": "",
            "reasoning": ""
        }

        sections = {
            "**facts**": "facts",
            "**issues**": "issues",
            "**petitionerArguments**": "petitioner_arguments",
            "**respondentArguments**": "respondent_arguments",
            "**holding**": "holding",
            "**reasoning**": "reasoning"
        }

        for marker, field in sections.items():
            start_idx = full_response.find(marker)
            if start_idx != -1:
                start_idx += len(marker)
                next_marker_positions = []
                for next_marker in sections.keys():
                    if next_marker != marker:
                        pos = full_response.find(next_marker, start_idx)
                        if pos != -1:
                            next_marker_positions.append(pos)

                if next_marker_positions:
                    end_idx = min(next_marker_positions)
                    section_content = full_response[start_idx:end_idx].strip()
                else:
                    section_content = full_response[start_idx:].strip()

                if section_content:
                    analysis[field] = section_content

        if all(value == "" for value in analysis.values()):
            for base_marker, field in sections.items():
                marker = base_marker.replace("**", "")
                alt_markers = [f"**{marker}**", f"*{marker}*", marker]

                for alt_marker in alt_markers:
                    if alt_marker in full_response:
                        parts = full_response.split(alt_marker, 1)
                        if len(parts) > 1:
                            content = parts[1]
                            for next_marker in sections.keys():
                                next_marker_clean = next_marker.replace("**", "")
                                alt_next_markers = [f"**{next_marker_clean}**", f"*{next_marker_clean}*", next_marker_clean]

                                for alt_next in alt_next_markers:
                                    if alt_next in content:
                                        content = content.split(alt_next)[0]
                                        break

                            analysis[field] = content.strip()
                            break

        if all(value == "" for value in analysis.values()):
            lines = full_response.split('\n')
            current_section = None

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                for marker, field in sections.items():
                    clean_marker = marker.replace("**", "").lower()
                    if clean_marker in line.lower() and len(line) < len(clean_marker) + 5:
                        current_section = field
                        break

                if current_section and not any(clean_marker in line.lower() for clean_marker in
                                              [m.replace("**", "").lower() for m in sections.keys()]):
                    analysis[current_section] += line + " "

        if all(value == "" for value in analysis.values()):
            paragraphs = [p for p in full_response.split('\n\n') if p.strip()]
            num_sections = len(sections)
            paragraphs_per_section = max(1, len(paragraphs) // num_sections)

            for i, field in enumerate(sections.values()):
                start_idx = i * paragraphs_per_section
                end_idx = (i + 1) * paragraphs_per_section
                if i == len(sections) - 1:
                    end_idx = len(paragraphs)

                if start_idx < len(paragraphs):
                    analysis[field] = " ".join(paragraphs[start_idx:end_idx])

        for field in analysis:
            for marker in sections.keys():
                if marker in analysis[field]:
                    analysis[field] = analysis[field].replace(marker, "")
            analysis[field] = analysis[field].strip()
            if not analysis[field]:
                analysis[field] = f"Information about {field.replace('_', ' ')} not found in document."

        return analysis

    except Exception as e:
        print(f"Legal analysis error: {str(e)}")
        return {
            "facts": f"Analysis failed: {str(e)}",
            "issues": "Unable to extract issues.",
            "petitioner_arguments": "Unable to extract petitioner arguments.",
            "respondent_arguments": "Unable to extract respondent arguments.",
            "holding": "Unable to extract holding.",
            "reasoning": "Unable to extract reasoning."
        }

@app.route('/upload', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        try:
            uploaded_file = request.files.get('document')
            if not uploaded_file:
                return jsonify({'status': 'Error', 'message': 'No file uploaded'})

            file_path = os.path.join(app.config['UPLOAD_FOLDER'], uploaded_file.filename)
            uploaded_file.save(file_path)

            print(f"File uploaded: {uploaded_file.filename}, Size: {os.path.getsize(file_path)} bytes")

            text = extract_text(file_path)
            print(f"Extracted text length: {len(text)} characters")
            if len(text) < 50:
                return jsonify({
                    'status': 'Error',
                    'message': 'Could not extract sufficient text from document'
                })

            if not db_connected or conn is None:
                if not connect_to_db():
                    return jsonify({
                        'status': 'Error',
                        'message': 'Database connection unavailable'
                    })

            print("Starting legal document analysis...")
            analysis = analyze_legal_document(text)
            print("Analysis completed")

            session = Session()
            try:
                new_doc = Document(
                    filename=uploaded_file.filename,
                    text=text,
                    facts=analysis['facts'],
                    issues=analysis['issues'],
                    petitioner_arguments=analysis['petitioner_arguments'],
                    respondent_arguments=analysis['respondent_arguments'],
                    holding=analysis['holding'],
                    reasoning=analysis['reasoning']
                )
                session.add(new_doc)
                session.commit()
                doc_id = new_doc.id

                result = {
                    'document_id': doc_id,
                    'status': 'Success',
                    'analysis': analysis
                }
                print("Returning response:", result)
                return jsonify(result)
            except Exception as e:
                session.rollback()
                raise e
            finally:
                session.close()
        except Exception as e:
            return jsonify({
                'status': 'Error',
                'message': f'Processing error: {str(e)}'
            })

    return render_template('index.html')


@app.route('/document/<int:doc_id>', methods=['GET'])
def get_document(doc_id):
    try:
        if not db_connected or conn is None:
            if not connect_to_db():
                return jsonify({
                    'status': 'Error',
                    'message': 'Database connection unavailable'
                })

        session = Session()
        try:
            doc = session.query(Document).filter(Document.id == doc_id).first()

            if not doc:
                return jsonify({
                    'status': 'Error',
                    'message': 'Document not found'
                })

            return jsonify({
                'status': 'Success',
                'document': {
                    'id': doc.id,
                    'filename': doc.filename,
                    'facts': doc.facts,
                    'issues': doc.issues,
                    'petitioner_arguments': doc.petitioner_arguments,
                    'respondent_arguments': doc.respondent_arguments,
                    'holding': doc.holding,
                    'reasoning': doc.reasoning
                }
            })
        except Exception as e:
            raise e
        finally:
            session.close()
    except Exception as e:
        return jsonify({
            'status': 'Error',
            'message': f'Error retrieving document: {str(e)}'
        })


def extract_result_manually(response_text, default_similarity=0.0):
    """Helper function to extract results from OpenAI response when JSON parsing fails"""
    result = {
        "is_relevant": "yes" in response_text.lower(),
        "similarity_score": default_similarity,
        "explanation": "No explanation provided.",
        "court_type": "Not specified",
        "year": "Unknown",
        "cites": [],
        "cited_by": [],
        "overruled_by": []
    }

    # Try to extract similarity score
    similarity_match = re.search(r'similarity_score"?\s*:\s*(0\.\d+)', response_text)
    if similarity_match:
        result["similarity_score"] = float(similarity_match.group(1))

    # Try to extract explanation
    explanation_match = re.search(r'explanation"?\s*:\s*"([^"]+)"', response_text)
    if explanation_match:
        result["explanation"] = explanation_match.group(1)

    # Try to determine court type from text
    if "supreme court" in response_text.lower():
        result["court_type"] = "Supreme Court"
    elif "high court" in response_text.lower():
        result["court_type"] = "High Court"

    # Try to find a year pattern in the text
    year_pattern = re.search(r'\b(19\d\d|20\d\d)\b', response_text)
    if year_pattern:
        result["year"] = year_pattern.group(1)

    # Extract some basic explanation
    lines = response_text.split('\n')
    for line in lines:
        if "3." in line or "explain" in line.lower():
            explanation = line.split(':', 1)[1].strip() if ':' in line else line.strip()
            result["explanation"] = explanation
            break

    # Try to extract citation lists
    cites_match = re.search(r'cites"?\s*:\s*\[(.*?)\]', response_text, re.DOTALL)
    if cites_match:
        cites_str = cites_match.group(1)
        result["cites"] = [cite.strip(' "\'') for cite in cites_str.split(',') if cite.strip()]

    cited_by_match = re.search(r'cited_by"?\s*:\s*\[(.*?)\]', response_text, re.DOTALL)
    if cited_by_match:
        cited_by_str = cited_by_match.group(1)
        result["cited_by"] = [cite.strip(' "\'') for cite in cited_by_str.split(',') if cite.strip()]

    overruled_match = re.search(r'overruled_by"?\s*:\s*\[(.*?)\]', response_text, re.DOTALL)
    if overruled_match:
        overruled_str = overruled_match.group(1)
        result["overruled_by"] = [cite.strip(' "\'') for cite in overruled_str.split(',') if cite.strip()]

    return result

def extract_year_from_text(initial_year, article_text):
    """Extract year from text using various patterns"""
    if initial_year != "Unknown" and re.match(r'\d{4}', str(initial_year)):
        return initial_year

    # Try to find year patterns in the text
    year_patterns = [
        r'Act of (\d{4})',
        r'Act, (\d{4})',
        r'dated (\d{4})',
        r'judgment of (\d{4})',
        r'decision in (\d{4})',
        r'(\d{4}) Act',
        r'(\d{4})[\s\)]Constitution',
        r'Constitution[\s\(](\d{4})'
    ]

    # Check for each pattern
    years_found = []
    for pattern in year_patterns:
        matches = re.findall(pattern, article_text)
        years_found.extend([int(y) for y in matches if 1700 < int(y) < 2026])

    # Also find any years in format 1900-2025
    general_years = re.findall(r'\b(19\d\d|20[0-2]\d)\b', article_text)
    years_found.extend([int(y) for y in general_years if 1700 < int(y) < 2026])

    # If we found years, use the most appropriate one
    if years_found:
        # Sort years and use the earliest one in reasonable range
        years_found.sort()
        for year in years_found:
            if 1950 <= year <= 2025:  # Reasonable range for Indian legal documents
                return year
        # If no year in reasonable range, just use the earliest
        if years_found:
            return years_found[0]

    return "Unknown"

def generate_embedding(text):
    """Generate embedding vector for text using OpenAI embedding model"""
    try:
        # Truncate text if too long
        max_length = 8000  # Adjust based on your embedding model's limits
        truncated_text = text[:max_length] if len(text) > max_length else text

        # Call OpenAI API to generate embedding
        response = client.embeddings.create(
            model="text-embedding-3-small",  # Use newer, faster model
            input=truncated_text
        )

        # Extract embedding from response
        embedding = response.data[0].embedding
        return embedding
    except Exception as e:
        print(f"Error generating embedding: {str(e)}")
        # Fallback to a random embedding for testing (should be replaced in production)
        import numpy as np
        return list(np.random.rand(1536))  # Simulate embedding vector

def find_relevant_documents_with_embeddings(document_text, max_articles=200):
    try:
        temp_dir = tempfile.mkdtemp(prefix="legal_docs_")
        print(f"Created temporary directory: {temp_dir}")

        document_embedding = generate_embedding(document_text)
        if document_embedding is None:
            print("Failed to generate embedding")
            shutil.rmtree(temp_dir)
            return []

        session = get_db_session()
        if not session:
            print("Database session error")
            shutil.rmtree(temp_dir)
            return []

        try:
            doc_embedding_array = np.array(document_embedding)
            chunks = session.query(ArticleChunk).filter(ArticleChunk.embedding.isnot(None)).limit(5000).all()
            print(f"Found {len(chunks)} chunks")

            chunk_similarities = []
            batch_size = 200
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i+batch_size]
                embeddings = np.array([chunk.embedding for chunk in batch])
                doc_reshaped = doc_embedding_array.reshape(1, -1)
                sims = cosine_similarity(doc_reshaped, embeddings)[0]
                for j, chunk in enumerate(batch):
                    chunk_similarities.append({
                        'chunk_id': chunk.id,
                        'article_id': chunk.article_id,
                        'chunk_number': chunk.chunk_number,
                        'similarity': float(sims[j]),
                        'chunk_text': chunk.chunk_text
                    })

            chunk_similarities.sort(key=lambda x: x['similarity'], reverse=True)
            top_chunks = chunk_similarities[:max_articles * 2]
            article_ids = {chunk['article_id'] for chunk in top_chunks}
            articles = session.query(Article2).filter(Article2.id.in_(article_ids)).all()
            article_dict = {a.id: a for a in articles}

            for chunk in top_chunks:
                article_id = chunk['article_id']
                if article_id in article_dict:
                    chunk['article_name'] = article_dict[article_id].name
                    full_text = ""
                    article_chunks = session.query(ArticleChunk).filter(
                        ArticleChunk.article_id == article_id
                    ).order_by(ArticleChunk.chunk_number).all()
                    for ac in article_chunks:
                        full_text += ac.chunk_text + " "
                    path = os.path.join(temp_dir, f"article_{article_id}.txt")
                    with open(path, 'w', encoding='utf-8') as f:
                        f.write(f"ARTICLE ID: {article_id}\nARTICLE NAME: {chunk['article_name']}\nCONTENT:\n{full_text}")
                    chunk['file_path'] = path
                    chunk['full_text'] = full_text

            results = []
            processed_ids = set()

            def process_article(chunk):
                article_id = chunk['article_id']
                if article_id not in article_dict:
                    return None

                file_path = chunk.get('file_path')
                if not file_path or not os.path.exists(file_path):
                    return None

                with open(file_path, 'r', encoding='utf-8') as f:
                    file_content = f.read()

                article_text = file_content.split("CONTENT:\n", 1)[1] if "CONTENT:\n" in file_content else file_content
                max_len = 4000
                doc_text_trimmed = document_text[:max_len]
                article_text_trimmed = article_text[:max_len]

                try:
                    prompt = f"""
You are a senior legal analyst specializing in Indian case law. Evaluate the following user document against a legal article from our database.
Give priority to articles from the **Supreme Court or High Court**, and treat any overlap of laws, citations, or judgments as significant. If there's even partial relevance due to shared law, citation, or issue, mark it as relevant.

# USER DOCUMENT:
{doc_text_trimmed}

# DATABASE ARTICLE: {chunk['article_name']}
{article_text_trimmed}

Perform the analysis and answer strictly in this format:
{{
  "is_relevant": true or false,
  "similarity_score": {chunk['similarity']:.4f},
  "explanation": "Why this article is or isn't relevant",
  "court_type": "Supreme Court or High Court or Not specified",
  "year": "e.g., 2005 or Unknown",
  "cites": ["Cited Case 1", "Cited Case 2"],
  "cited_by": ["Referred Case A", "Referred Case B"],
  "overruled_by": ["Overruled Case X"]
}}
"""

                    response = client.chat.completions.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": "You are a legal document analysis expert for Indian law."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=800,
                        temperature=0.3
                    )

                    response_text = response.choices[0].message.content
                    match = re.search(r'\{.*\}', response_text, re.DOTALL)
                    if match:
                        json_str = match.group(0)
                        result = json.loads(json_str)
                    else:
                        result = extract_result_manually(response_text, chunk['similarity'])

                    if result.get("similarity_score", 0) <= 0:
                        result["similarity_score"] = chunk['similarity']
                    if result.get("similarity_score", 0) > 0.1:
                        result["is_relevant"] = True

                    if result.get("similarity_score", 0) <= 0.05 and not result.get("is_relevant"):
                        return None

                    year = extract_year_from_text(result.get("year", "Unknown"), article_text)

                    return {
                        'id': article_id,
                        'name': chunk['article_name'],
                        'similarity': result["similarity_score"],
                        'explanation': result.get("explanation", ""),
                        'type': 'relevant',
                        'court_type': result.get("court_type", "Not specified"),
                        'year': year,
                        'cited_by': result.get("cited_by", []),
                        'cites': result.get("cites", []),
                        'overruled_by': result.get("overruled_by", []),
                        'snippet': f"<h2 class='doc_title'>{chunk['article_name']}</h2>{chunk['chunk_text'][:300]}..."
                    }

                except Exception as e:
                    print(f"Error with article {article_id}: {str(e)}")
                    return None

            unique_chunks = []
            for chunk in top_chunks:
                if chunk['article_id'] not in processed_ids:
                    processed_ids.add(chunk['article_id'])
                    unique_chunks.append(chunk)

            with ThreadPoolExecutor(max_workers=5) as executor:
                article_results = list(executor.map(process_article, unique_chunks))

            for res in article_results:
                if res:
                    results.append(res)

            results.sort(key=lambda x: x['similarity'], reverse=True)

            if len(results) < 10:
                print("Fewer than 10 relevant documents found. Including additional borderline cases...")
                extra_results = [r for r in article_results if r and r not in results and r['similarity'] > 0.05]
                results.extend(extra_results[:10 - len(results)])

            results = results[:max_articles]
            print(f"Final relevant results count: {len(results)}")

            shutil.rmtree(temp_dir)
            session.commit()
            return results

        except Exception as e:
            print(f"Error during session work: {str(e)}")
            shutil.rmtree(temp_dir)
            return []
        finally:
            session.close()

    except Exception as e:
        print(f"Critical error: {str(e)}")
        return []


@app.route('/find-relevant-ai/<int:doc_id>', methods=['GET'])
def find_relevant_ai(doc_id):
    try:
        print(f"Starting find-relevant-ai for document ID: {doc_id}")

        # Get database session with retry logic
        session = get_db_session()
        if session is None:
            return jsonify({
                'status': 'Error',
                'message': 'Database connection unavailable'
            })

        try:
            # Fetch document
            print(f"Fetching document with ID: {doc_id}")
            doc = session.query(Document).filter(Document.id == doc_id).first()

            if not doc:
                print(f"Document with ID {doc_id} not found")
                return jsonify({
                    'status': 'Error',
                    'message': 'Document not found'
                })

            print(f"Document found, text length: {len(doc.text) if hasattr(doc, 'text') else 'No text'}")

            # Find relevant documents using embeddings
            print("Finding relevant documents using embeddings")
            try:
                relevant_docs = find_relevant_documents_with_embeddings(doc.text, max_articles=200)

                # Transform results into nodes and links format
                nodes = []
                links = []
                node_ids = {}  # Track which nodes are already added

                # Create a citation map for generating links
                citation_map = {}
                for doc_info in relevant_docs:
                    doc_id = doc_info['id']
                    doc_name = doc_info['name']
                    citation_map[doc_id] = {
                        'id': doc_id,
                        'name': doc_name,
                        'cites': doc_info['cites'],
                        'cited_by': doc_info['cited_by'],
                        'overruled_by': doc_info['overruled_by']
                    }

                # First, add all relevant documents as nodes
                for doc_info in relevant_docs:
                    # Skip if we've already added this node
                    if doc_info['id'] in node_ids:
                        continue

                    # Add the node
                    nodes.append({
                        'id': doc_info['id'],
                        'name': doc_info['name'],
                        'type': 'relevant',
                        'court_type': doc_info['court_type'],
                        'year': doc_info['year'],
                        'snippet': doc_info['snippet'],
                        'explanation': doc_info['explanation'],
                        'similarity_score': doc_info['similarity']
                    })

                    # Mark this node as added
                    node_ids[doc_info['id']] = True

                # Create a mapping from document name to ID for link resolution
                name_to_id = {node['name']: node['id'] for node in nodes}

                # IMPROVED LINK CREATION - Match by partial name when necessary
                def find_node_id_by_name(name):
                    # Direct match
                    if name in name_to_id:
                        return name_to_id[name]

                    # Try partial matching
                    for node_name, node_id in name_to_id.items():
                        # Check if this is a substring of a longer name
                        if name in node_name:
                            return node_id

                        # Check if name parts match
                        name_parts = name.lower().split()
                        node_parts = node_name.lower().split()

                        # If at least 3 unique words match, it's likely the same document
                        if len(name_parts) >= 3 and len(node_parts) >= 3:
                            matching_words = set(name_parts) & set(node_parts)
                            if len(matching_words) >= 3:
                                return node_id

                    return None

                # Then create links between documents
                processed_links = set()  # Track unique links

                for doc_info in relevant_docs:
                    source_id = doc_info['id']

                    # Add citation links
                    for cited_doc in doc_info['cites']:
                        target_id = find_node_id_by_name(cited_doc)
                        if target_id and target_id != source_id:
                            link_key = f"{source_id}-{target_id}-cites"
                            if link_key not in processed_links:
                                links.append({
                                    'source': source_id,
                                    'target': target_id,
                                    'relation': 'cites'
                                })
                                processed_links.add(link_key)

                    # Add cited_by links
                    for citing_doc in doc_info['cited_by']:
                        target_id = find_node_id_by_name(citing_doc)
                        if target_id and target_id != source_id:
                            link_key = f"{target_id}-{source_id}-citedBy"
                            if link_key not in processed_links:
                                links.append({
                                    'source': target_id,
                                    'target': source_id,
                                    'relation': 'citedBy'
                                })
                                processed_links.add(link_key)

                    # Add overruled_by links
                    for overruling_doc in doc_info['overruled_by']:
                        target_id = find_node_id_by_name(overruling_doc)
                        if target_id and target_id != source_id:
                            link_key = f"{source_id}-{target_id}-overruledBy"
                            if link_key not in processed_links:
                                links.append({
                                    'source': source_id,
                                    'target': target_id,
                                    'relation': 'overruledBy'
                                })
                                processed_links.add(link_key)

                # If we have fewer than 5 real links but multiple nodes, create some additional connections
                if len(links) < 5 and len(nodes) >= 2:
                    # Create reasonable connections based on document types and years
                    for i in range(len(nodes)):
                        for j in range(len(nodes)):
                            if i != j and len(links) < max(10, len(nodes)):
                                # Don't add too many artificial links
                                source_node = nodes[i]
                                target_node = nodes[j]

                                # Skip if we already have a link between these nodes
                                link_exists = any(
                                    (link['source'] == source_node['id'] and link['target'] == target_node['id']) or
                                    (link['source'] == target_node['id'] and link['target'] == source_node['id'])
                                    for link in links
                                )

                                if not link_exists:
                                    # Determine a reasonable relation based on years
                                    relation = 'cites'
                                    source_year = source_node.get('year', 'Unknown')
                                    target_year = target_node.get('year', 'Unknown')

                                    # Try to convert years to integers
                                    try:
                                        source_year_int = int(source_year) if source_year != 'Unknown' else 0
                                        target_year_int = int(target_year) if target_year != 'Unknown' else 0

                                        # Newer documents cite older ones
                                        if source_year_int > target_year_int and target_year_int > 0:
                                            relation = 'cites'
                                        # Older documents are cited by newer ones
                                        elif target_year_int > source_year_int and source_year_int > 0:
                                            relation = 'citedBy'
                                        # Same year or unknown, assign randomly
                                        else:
                                            relation = random.choice(['cites', 'citedBy', 'overruledBy'])
                                    except (ValueError, TypeError):
                                        # If years can't be compared, use random relation
                                        relation = random.choice(['cites', 'citedBy', 'overruledBy'])

                                    # Create the link
                                    links.append({
                                        'source': source_node['id'],
                                        'target': target_node['id'],
                                        'relation': relation
                                    })

                print(f"Returning {len(nodes)} nodes and {len(links)} links")
                return jsonify({
                    'document_id': doc_id,
                    'citation_legend': {
                        'incoming': 'Cited by - Other cases that reference this document',
                        'outgoing': 'Cites - Cases that this document references',
                        'overruled_by': 'Overruled By - Cases that have overturned this document\'s ruling',
                        'year': 'Year - The year when this judgment was delivered'
                    },
                    'nodes': nodes,
                    'links': links
                })
            except Exception as e:
                print(f"Error in find_relevant_documents_with_embeddings: {str(e)}")
                # Return error response
                return jsonify({
                    'status': 'Error',
                    'message': f'Error finding relevant documents: {str(e)}'
                })

        except Exception as e:
            print(f"Session error: {str(e)}")
            raise e
        finally:
            session.close()
            print("Session closed")
    except Exception as e:
        print(f"Overall error in find_relevant_ai: {str(e)}")
        return jsonify({
            'status': 'Error',
            'message': f'Error finding relevant documents: {str(e)}'
        })



######################## new code ########################



def get_text_snippet(text, search_term, context_length=150):
    """Extract text snippet showing the search term in context"""
    search_term = search_term.lower()
    text_lower = text.lower()

    index = text_lower.find(search_term)
    if index == -1:
        return text[:300] + "..." if len(text) > 300 else text

    start = max(0, index - context_length)
    end = min(len(text), index + len(search_term) + context_length)

    snippet = ""
    if start > 0:
        snippet += "..."

    snippet += text[start:end]

    if end < len(text):
        snippet += "..."

    return snippet

@app.route('/search-articles', methods=['POST'])
def search_articles():
    try:
        # Get search text from request
        data = request.get_json()
        if not data or 'search_text' not in data:
            return jsonify({
                'status': 'Error',
                'message': 'Search text is required'
            }), 400

        search_text = data['search_text']

        # Verify database connection
        if not db_connected or conn is None:
            if not connect_to_db():
                return jsonify({
                    'status': 'Error',
                    'message': 'Database connection unavailable'
                }), 500

        # Search for articles containing the search text
        session = Session()
        try:
            limit = data.get('limit', 10)  # Default to 10 results
            offset = data.get('offset', 0)  # Default to first page

            # Search both in name and text columns
            articles = session.query(Article2).filter(
                (Article2.name.ilike(f'%{search_text}%')) |
                (Article2.text.ilike(f'%{search_text}%'))
            ).limit(limit).offset(offset).all()

            # Count total matches for pagination
            total_count = session.query(Article2).filter(
                (Article2.name.ilike(f'%{search_text}%')) |
                (Article2.text.ilike(f'%{search_text}%'))
            ).count()

            results = []
            for article in articles:
                # Get text snippet around the matched content
                text_snippet = get_text_snippet(article.text, search_text)

                results.append({
                    'id': article.id,
                    'name': article.name,
                    'snippet': text_snippet,
                    'match_type': 'Direct keyword match'
                })

            return jsonify({
                'status': 'Success',
                'search_text': search_text,
                'results': results,
                'total_matches': total_count,
                'limit': limit,
                'offset': offset
            })

        except Exception as e:
            session.rollback()
            raise e
        finally:
            session.close()

    except Exception as e:
        return jsonify({
            'status': 'Error',
            'message': f'Search error: {str(e)}'
        }), 500


@app.route('/api/conversations', methods=['GET'])
@token_required
def get_all_conversations(user_id):
    try:
        # Get pagination parameters
        page = request.args.get('page', default=1, type=int)
        per_page = request.args.get('per_page', default=10, type=int)

        # Calculate offset
        offset = (page - 1) * per_page

        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Database connection failed'}), 500

        cur = conn.cursor(cursor_factory=RealDictCursor)

        # First get distinct conversation_ids for the user with latest message
        cur.execute('''
            SELECT DISTINCT ON (conversation_id)
                conversation_id,
                created_at,
                message,
                response
            FROM amicus
            WHERE user_id = %s
            ORDER BY conversation_id, created_at DESC
            LIMIT %s OFFSET %s
        ''', (str(user_id), per_page, offset))

        conversations = cur.fetchall()

        # Get total count for pagination
        cur.execute('SELECT COUNT(DISTINCT conversation_id) FROM amicus WHERE user_id = %s', (str(user_id),))
        total_count = cur.fetchone()['count']

        # For each conversation, get message count
        for conv in conversations:
            cur.execute('''
                SELECT COUNT(*) FROM amicus
                WHERE conversation_id = %s AND user_id = %s
            ''', (conv['conversation_id'], str(user_id)))

            conv['message_count'] = cur.fetchone()['count']
            # Format the timestamp
            if 'created_at' in conv and conv['created_at']:
                conv['created_at'] = conv['created_at'].isoformat()

        cur.close()
        conn.close()

        return jsonify({
            'conversations': conversations,
            'pagination': {
                'total': total_count,
                'page': page,
                'per_page': per_page,
                'total_pages': math.ceil(total_count / per_page)
            }
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/supreme-court', methods=['GET'])
def get_supreme_court_data():
    try:
        # Get database session with retry logic
        session = get_db_session()
        if session is None:
            return jsonify({
                'status': 'Error',
                'message': 'Database connection unavailable'
            }), 500

        try:
            # Fetch Supreme Court data
            supreme_court_data = session.query(Article2).filter(
                Article2.court_type == CourtType.SUPREME_COURT
            ).all()

            # Format the data
            result = []
            for article in supreme_court_data:
                result.append({
                    'id': article.id,
                    'name': article.name,
                    'year': article.year,
                    'court_type': article.court_type.value,
                    # Optionally add more fields as needed
                })

            return jsonify({
                'status': 'Success',
                'count': len(result),
                'data': result
            })

        except Exception as e:
            return jsonify({
                'status': 'Error',
                'message': f'Error fetching Supreme Court data: {str(e)}'
            }), 500
        finally:
            session.close()
    except Exception as e:
        return jsonify({
            'status': 'Error',
            'message': f'Server error: {str(e)}'
        }), 500

# API 2: Fetch High Court data
@app.route('/api/high-court', methods=['GET'])
def get_high_court_data():
    try:
        # Get database session with retry logic
        session = get_db_session()
        if session is None:
            return jsonify({
                'status': 'Error',
                'message': 'Database connection unavailable'
            }), 500

        try:
            # Fetch High Court data
            high_court_data = session.query(Article2).filter(
                Article2.court_type == CourtType.HIGH_COURT
            ).all()

            # Format the data
            result = []
            for article in high_court_data:
                result.append({
                    'id': article.id,
                    'name': article.name,
                    'year': article.year,
                    'court_type': article.court_type.value,
                    # Optionally add more fields as needed
                })

            return jsonify({
                'status': 'Success',
                'count': len(result),
                'data': result
            })

        except Exception as e:
            return jsonify({
                'status': 'Error',
                'message': f'Error fetching High Court data: {str(e)}'
            }), 500
        finally:
            session.close()
    except Exception as e:
        return jsonify({
            'status': 'Error',
            'message': f'Server error: {str(e)}'
        }), 500

####################### column api #########################

@app.route('/api/columns', methods=['POST'])
@token_required
def create_column(user_id):
    try:
        # Check if request is JSON or form data
        data = {}
        image_data = None

        if request.is_json:
            data = request.json
            if 'image' in data and data['image'] and data['image'].startswith('data:image/'):
                image_data = data['image']
        else:
            data = request.form.to_dict()
            if 'image' in request.files and request.files['image'].filename:
                file = request.files['image']
                if file and allowed_file(file.filename):
                    image_data = file

        # Validate required fields
        if 'title' not in data or not data['title']:
            return jsonify({'error': 'Title is required'}), 400

        if 'content' not in data or not data['content']:
            return jsonify({'error': 'Content is required'}), 400

        # Process and save the image if provided
        image_path = None

        if image_data:
            # Handle base64 image
            if isinstance(image_data, str) and image_data.startswith('data:image/'):
                image_path = save_base64_image(image_data, user_id)
            # Handle file upload
            elif hasattr(image_data, 'filename'):
                filename = secure_filename(f"column_{user_id}_{uuid.uuid4().hex}_{image_data.filename}")
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                image_data.save(filepath)
                image_path = filepath

        # Save column to database
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        cur.execute('''
            INSERT INTO columns (title, content, image_path, user_id)
            VALUES (%s, %s, %s, %s)
            RETURNING id, title, content, image_path, user_id, created_at, updated_at
        ''', (
            data['title'],
            data['content'],
            image_path,  # This can now be None
            user_id
        ))

        new_column = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()

        # Format response
        result = dict(new_column)
        if image_path:
            result['image_url'] = f"/api/images/{os.path.basename(image_path)}"
        else:
            result['image_url'] = None

        return jsonify({
            'message': 'Column created successfully',
            'column': result
        }), 201

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/columns', methods=['GET'])
def get_columns():
    try:
        # Get pagination parameters
        page = request.args.get('page', default=1, type=int)
        per_page = request.args.get('per_page', default=10, type=int)

        # Calculate offset
        offset = (page - 1) * per_page

        # Get user_id filter if provided
        user_id_filter = request.args.get('user_id')

        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        # Build query based on whether user_id filter is provided
        if user_id_filter:
            cur.execute('''
                SELECT id, title, content, image_path, user_id, created_at, updated_at
                FROM columns
                WHERE user_id = %s
                ORDER BY created_at DESC
                LIMIT %s OFFSET %s
            ''', (user_id_filter, per_page, offset))
        else:
            cur.execute('''
                SELECT id, title, content, image_path, user_id, created_at, updated_at
                FROM columns
                ORDER BY created_at DESC
                LIMIT %s OFFSET %s
            ''', (per_page, offset))

        columns = cur.fetchall()

        # Get total count for pagination
        if user_id_filter:
            cur.execute('SELECT COUNT(*) FROM columns WHERE user_id = %s', (user_id_filter,))
        else:
            cur.execute('SELECT COUNT(*) FROM columns')

        total_count = cur.fetchone()['count']

        cur.close()
        conn.close()

        # Format response with image URLs (checking for None)
        for column in columns:
            if column['image_path']:
                column['image_url'] = f"/api/images/{os.path.basename(column['image_path'])}"
            else:
                column['image_url'] = None

        return jsonify({
            'columns': columns,
            'pagination': {
                'total': total_count,
                'page': page,
                'per_page': per_page,
                'total_pages': math.ceil(total_count / per_page)
            }
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/columns/<int:column_id>', methods=['GET'])
def get_column(column_id):
    try:
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        cur.execute('''
            SELECT id, title, content, image_path, user_id, created_at, updated_at
            FROM columns
            WHERE id = %s
        ''', (column_id,))

        column = cur.fetchone()

        cur.close()
        conn.close()

        if not column:
            return jsonify({'error': 'Column not found'}), 404

        # Format response with image URL (with None check)
        if column['image_path']:
            column['image_url'] = f"/api/images/{os.path.basename(column['image_path'])}"
        else:
            column['image_url'] = None

        return jsonify(column), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/columns/<int:column_id>', methods=['PUT'])
@token_required
def update_column(user_id, column_id):
    try:
        # Check if request is JSON or form data
        data = {}
        image_data = None

        if request.is_json:
            data = request.json
            if 'image' in data and data['image'] and data['image'].startswith('data:image/'):
                image_data = data['image']
        else:
            data = request.form.to_dict()
            if 'image' in request.files and request.files['image'].filename:
                file = request.files['image']
                if file and allowed_file(file.filename):
                    image_data = file

        # Validate required fields
        if ('title' not in data or not data['title']) and ('content' not in data or not data['content']):
            return jsonify({'error': 'Title or content must be provided'}), 400

        # First check if column exists and belongs to user
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)

        cur.execute('''
            SELECT id, image_path, user_id
            FROM columns
            WHERE id = %s
        ''', (column_id,))

        column = cur.fetchone()

        if not column:
            cur.close()
            conn.close()
            return jsonify({'error': 'Column not found'}), 404

        if column['user_id'] != user_id:
            cur.close()
            conn.close()
            return jsonify({'error': 'Unauthorized to modify this column'}), 403

        # Process image if provided
        image_path = column['image_path']  # Default to existing path

        if image_data:
            # Handle base64 image
            if isinstance(image_data, str) and image_data.startswith('data:image/'):
                image_path = save_base64_image(image_data, user_id)
            # Handle file upload
            elif hasattr(image_data, 'filename'):
                filename = secure_filename(f"column_{user_id}_{uuid.uuid4().hex}_{image_data.filename}")
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                image_data.save(filepath)
                image_path = filepath

        # Build update query based on what's provided
        update_fields = []
        update_values = []

        if 'title' in data and data['title']:
            update_fields.append("title = %s")
            update_values.append(data['title'])

        if 'content' in data and data['content']:
            update_fields.append("content = %s")
            update_values.append(data['content'])

        if image_path != column['image_path']:
            update_fields.append("image_path = %s")
            update_values.append(image_path)

        update_fields.append("updated_at = NOW()")

        # Add column_id at the end for WHERE clause
        update_values.append(column_id)

        # Execute update
        cur.execute(f'''
            UPDATE columns
            SET {', '.join(update_fields)}
            WHERE id = %s
            RETURNING id, title, content, image_path, user_id, created_at, updated_at
        ''', tuple(update_values))

        updated_column = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()

        # Format response with image URL (with None check)
        result = dict(updated_column)
        if updated_column['image_path']:
            result['image_url'] = f"/api/images/{os.path.basename(updated_column['image_path'])}"
        else:
            result['image_url'] = None

        return jsonify({
            'message': 'Column updated successfully',
            'column': result
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500



@app.route('/', methods=['GET'])
def home():
    return jsonify({'message': 'Welcome to the API home page'}), 200


if __name__ == '__main__':
    app.run(debug=True, port=5000)
